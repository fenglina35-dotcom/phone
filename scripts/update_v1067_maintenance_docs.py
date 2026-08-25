from pathlib import Path

from docx import Document


ROOT = Path(__file__).resolve().parents[1]
DOCS = ROOT / "docs/maintenance"
MARKER = "v1067／1.0.190"


def append_section(filename, title, paragraphs):
    path = DOCS / filename
    document = Document(path)
    if any(MARKER in paragraph.text for paragraph in document.paragraphs):
        print(f"SKIP={path.name}")
        return
    document.add_heading(title, level=1)
    for paragraph in paragraphs:
        document.add_paragraph(paragraph)
    document.save(path)
    print(f"UPDATED={path.name}")


append_section(
    "AI开发项目_项目说明文档.docx",
    "v1067／1.0.190 私人 App Intl 热点隔离与白屏恢复（2026-08-25）",
    [
        "网页核心升级为 v1067，私人 iOS 升级为 1.0.190（190），原生桥继续为 25。本次依据真机 Instruments 调用栈修复私人 App 的 WebContent 主线程热点；不改角色模型、消息、通知、伴生指令和语音业务逻辑。",
        "关键证据：用户录制中约 96 秒的 WebContent 样本持续落在 JavaScriptCore 的 Intl.DateTimeFormat／ICU 构造链和 DOM 定时器路径。该连续热点可同时解释点击延迟、发热、WebContent 白屏终止，以及依赖页面 JavaScript 的消息同步暂时停止。v1066 只处理伴生日期入口，真机仍复现，不能算完成。",
        "v1067 在私人 App 启动时由 Swift 注入本机时区及可用时区偏移；私人 App 的设备时区、角色时间、时区校验、时区列表、数字与日期显示不再进入 Intl/toLocaleString 热路径。网页版保留原 Intl 能力和既有时区语义。",
        "新增 WKWebView WebContent 终止恢复：两分钟内最多自动重载两次，并沿用原有原生存档恢复；连续第三次停止自动重载并显示失败，防止无限热重启。自动恢复不清缓存、不伪造数据。",
        "Windows 全量自动回归通过；Mac 编译、签名以及 v1067 真机长时间验证仍待完成。安装后须验证冷启动、持续点击、后台消息、白屏恢复和温度变化，不能把 MacReady 安装包写成已在真机修好。",
    ],
)

append_section(
    "AI开发项目_Bug记录模板.docx",
    "v1067／1.0.190 私人 App 卡顿发热与白屏 Bug 记录（2026-08-25）",
    [
        "实际现象：私人 App 偶发严重卡顿，滑动可能仍可进行，但点击响应很慢；打开小应用偶发白屏并回主屏；手机温度快速上升；同期后台通知可能仍到系统，但页面消息同步和交互会暂停。状态也可能突然恢复，恢复后温度下降。",
        "真实证据：Instruments 录制显示 WebContent 主线程约 96 秒持续消耗在 Intl.DateTimeFormat／ICU 构造和 DOM 定时器链，构成可重复解释全部症状的共同根因。旧服务器恢复、手机硬件、单一图片、定位和 APNs 不能解释本机 WebContent 的该调用栈。",
        "已排除的错误方案：不能再只改 companionUsageDayAt、盲目放慢全部定时器、删除后台能力、改模型或消息协议；v1066 的局部 formatter 修复已被真机复现否定。也不能以系统进度、缓存数据或假角色回复掩盖页面停摆。",
        "本次修改：私人 App 的常用时间和数字格式化全部走无 Intl 快速路径，时区数据由原生一次注入；网页逻辑保留。WKWebView 被系统终止后只做有上限的安全重载，不无限循环。",
        "验证边界：Windows 自动测试可证明语法、私有路径不触发 Intl、Bundle 同步和核心回归边界；不能证明 Mac 编译、签名或真实 iPhone 已恢复。真机复现若仍存在，下一步必须使用同版本 v1067 的新 trace 对比热点，不得继续在旧假设上修改。",
    ],
)

append_section(
    "AI开发项目_Bug修改规范.docx",
    "v1067／1.0.190 私人 WebContent 性能修复规范（2026-08-25）",
    [
        "证据门槛：偶发卡顿发热必须以同一安装版本的 Time Profiler、Hangs 和 Thermal State 为准；调用栈、持续时间和用户复现时间必须能互相对应。没有新证据时禁止重复修改同一处。",
        "完整性门槛：修复某个 Intl 入口后，必须静态搜索和动态拦截私人 App 的其余 Intl/toLocaleString 热路径；只修一处但仍有同类入口时，不得称为根因已修复。网页版需要的国际化能力应单独保留。",
        "恢复门槛：WKWebView WebContent 终止可以重载，但必须限制次数、保留原生权威存档并禁止清数据；禁止无上限重载造成更严重发热和闪屏。",
        "保护边界：不得修改真实模型回复、朋友圈真实回复、APNs 与消息落库、远控真实字幕、伴生真实数据、通话、内外置语音配置、共同相册和私人/网页能力隔离。系统文字不得冒充 AI，旧快照不得冒充实时数据。",
        "发布门槛：版本号、网页壳、私人 Bundle、Xcode marketing/build、维护记录、自动测试和 ZIP 必须一致；使用显式文件列表提交，禁止 git add -A、git clean、reset --hard 和整版旧文件覆盖。MacReady 不等于 Mac 编译或真机验证。",
    ],
)

txt_path = DOCS / "AI开发项目_Bug记录模板.txt"
text = txt_path.read_text(encoding="utf-8")
if MARKER not in text:
    text += """

v1067／1.0.190 私人 App 卡顿发热与白屏（2026-08-25）
- 现象：偶发点击延迟、快速发热、白屏回主屏，并伴随页面消息同步暂停。
- 证据：真机 trace 中 WebContent 约 96 秒持续处于 Intl.DateTimeFormat／ICU 与 DOM 定时器链；v1066 的单入口修复已被真机复现否定。
- 修复：私人 App 常用时间和数字格式化隔离 Intl，由 Swift 注入时区；WebContent 终止只允许两分钟内安全重载两次。网页国际化与全部核心业务链路不改。
- 验证：Windows 自动回归通过；Mac 编译、签名和 v1067 真机长时间结果仍待验证。
"""
    txt_path.write_text(text, encoding="utf-8")
    print(f"UPDATED={txt_path.name}")
else:
    print(f"SKIP={txt_path.name}")

from pathlib import Path

from docx import Document


ROOT = Path(__file__).resolve().parents[1]
DOCS = ROOT / "docs" / "maintenance"


def append_section(path: Path, heading: str, paragraphs: list[str]) -> None:
    doc = Document(path)
    if any(p.text.strip() == heading for p in doc.paragraphs):
        raise RuntimeError(f"section already exists: {heading}")
    doc.add_heading(heading, level=1)
    for paragraph in paragraphs:
        doc.add_paragraph(paragraph)
    doc.save(path)


append_section(
    DOCS / "AI开发项目_项目说明文档.docx",
    "2026 09 06 私人线下原生聚焦修复候选",
    [
        "发布范围：私人内置网页 v1188／私人 iOS 1.0.314 (314)，原生桥 35；公开网页仍为 v1184，公开网页源码和微信输入框不改动。",
        "真机复核结论：v1187 新增的 offComposerPinLatest 在 off_in 的 touchstart 和 pointerdown 阶段重复写入 offbg.scrollTop。该时机与 WebKit 的点击命中、光标定位和键盘动画重叠，能够同时解释光标坐标再次错误、首次打开弹跳及收起时输入栏二段回落。",
        "候选修复：删除线下输入框触摸聚焦阶段的滚动改写，恢复 v1180 的原生聚焦路径；保留 16 像素输入字体、自动增高、触摸去重和旁白切换不重复 focus。没有照搬微信整页布局，没有修改全局 WKWebView、SwiftUI 键盘安全区或微信的 chat-inputbar 规则。",
        "验证边界：Windows 只能验证源码作用域、公开网页零改动、静态约束、测试和安装包内容。真实光标与动画必须在 Mac/Xcode 编译签名后由实际 iPhone 验证，未通过前仅为候选。",
    ],
)

append_section(
    DOCS / "AI开发项目_Bug记录模板.docx",
    "v1188 私人 iOS314 线下光标与收起二段回落记录",
    [
        "真机现象：v1187 已恢复微信，但共同生活和线下约会的光标坐标再次偏离；首次打开仍可能弹跳，收起键盘时键盘先下降，输入栏随后像移动格子一样回落并短暂停顿。",
        "根因：v1187 为解决旧消息停留，在 off_in 的 touchstart 与 pointerdown 捕获阶段都调用 offComposerPinLatest，同一次触摸可能两次写入内部消息区 scrollTop。写入发生在 WebKit 完成 caret hit testing 和键盘位移动画之前，形成额外滚动时间线。",
        "未采用方案：不把微信整套页面 CSS 和布局直接复制到线下。微信与线下的消息容器、旁白按钮、多人目标选择和输入栏结构不同；硬复制会扩大作用域，并可能重现 v1181 以后 fixed／absolute 与宿主缩放竞争。",
        "v1188 候选处理：彻底删除 offComposerPinLatest 及触摸、指针入口中的调用；保留与布局无关且已验证的 16 像素字体、自动增高、触摸去重和旁白切换不断焦。微信、公共网页、WKWebView 和 SwiftUI 根层零改动。",
        "当前状态：Windows 针对性与全量回归通过后可打包为 Mac 待编译源码。真实 iPhone 必须验证首次聚焦、光标点位、连续开关三轮、收起同步、旧消息状态、旁白互切和微信零回归。",
    ],
)

append_section(
    DOCS / "AI开发项目_Bug修改规范.docx",
    "新增规则 禁止在输入命中阶段改写滚动位置",
    [
        "iOS WKWebView 的 textarea 在 touchstart、pointerdown 和原生 caret hit testing 完成前，不得同步改写任何祖先消息区或外层页面的 scrollTop、contentOffset、transform、bottom 或高度。",
        "同一物理触摸可能依次产生 touchstart、pointerdown 和 click；不得让多个入口重复执行布局或滚动副作用。去重不只检查业务动作次数，还必须检查几何状态是否被重复写入。",
        "需要默认显示最后一条消息时，应在页面进入或消息渲染完成阶段处理，不得把滚动校正塞进输入框触摸阶段。若用户主动停在旧消息，是否跳到底部应作为独立交互规则验证。",
        "私人线下修复不得复制微信整页布局。只能复用已证明与 DOM 结构无关的原则；修改前要核对消息滚动容器、输入栏兄弟关系、旁白或目标控件、fixed 或 absolute 祖先及宿主键盘策略。",
    ],
)

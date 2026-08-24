from pathlib import Path

from docx import Document


ROOT = Path(__file__).resolve().parents[1]
DOCS = ROOT / "docs/maintenance"
MARKER = "v1060／1.0.183 微信忙碌与外卖澄清续接修复"


def append_section(filename, title, paragraphs):
    path = DOCS / filename
    document = Document(path)
    if any(MARKER in paragraph.text for paragraph in document.paragraphs):
        print(f"SKIP={path.name}")
        return
    document.add_heading(title, level=1)
    for paragraph in paragraphs:
        document.add_paragraph(paragraph)
    document.save(path)
    print(f"UPDATED={path.name}")


append_section(
    "AI开发项目_项目说明文档.docx",
    "v1060／1.0.183 微信忙碌与外卖澄清续接修复（2026-08-25）",
    [
        "网页核心升级为 v1060，私人 iOS 升级为 1.0.183（183），原生桥继续为 25。角色微信设置新增可手动开关的忙碌状态测试；忙碌期间保持在线并积累普通用户消息，关闭后通过原真实模型与完整人设合并回应，不写入固定角色台词。",
        "外卖起送价澄清增加受限续接：用户点名灌汤包或煎饺时只补对应候选；回答都要、两个都要或全加时，补齐角色在本轮明确列出的全部候选。续接保留原 taskId、门店、主商品和进度，并对同一消息幂等。",
        "自然语言补偿只允许作用于已有、未结束且账号会话匹配的 minimum_order 任务；普通聊天、旧消息、含糊答案和已结束任务不能创建新外卖任务。普通角色回复、朋友圈、后台通知、远控字幕、共同生活、转账和外置语音线路保持不变。",
        "根网页与私人 PhoneWeb.bundle 同步。Windows 完成语法、全量自动测试和 ZIP 结构检查；Mac 编译、签名和真实 iPhone 覆盖安装仍待执行。",
    ],
)

append_section(
    "AI开发项目_Bug记录模板.docx",
    "v1060／1.0.183 微信忙碌与外卖澄清续接 Bug 记录（2026-08-25）",
    [
        "现象：真实外卖因起送价不足暂停后，角色列出煎饺和灌汤包，用户回答都要，角色虽自然回复同意但模型未返回结构化外卖动作，自动化因此没有继续。",
        "根因：授权闸门只接收当前模型结构化动作；自然可见回复不能作为新任务授权。这一安全边界正确，但缺少对已有 minimum_order 澄清任务的窄范围答案解析。",
        "修复：只从当前用户消息之前的角色候选句中提取名称，点名候选只补该项，明确全部才补全；使用原 taskId 进入既有换品续接，并记录消息标识防止重复。含糊回答不猜测、不启动。",
        "验证：新增单选、全选、重复调用和含糊回答测试；全量回归覆盖外卖状态机、授权、普通角色回复、朋友圈、后台通知和私人语音隔离。",
    ],
)

append_section(
    "AI开发项目_Bug修改规范.docx",
    "v1060／1.0.183 微信忙碌与外卖澄清续接规范（2026-08-25）",
    [
        "忙碌规范：测试开关只延后普通用户消息的回复；功能事件仍走原可靠链路。关闭时必须使用真实模型按角色人设回应，生成失败保留待回复消息，禁止固定话术冒充角色。",
        "澄清规范：自然答案只能续接当前角色、账号、会话下已有且未结束的 minimum_order 任务。候选必须来自角色本轮真实可见列举；单选只取明确名称，全部选择必须有都要、两个都要、全加等明确表达。",
        "幂等规范：续接沿用原 taskId 和原主商品；同一 messageId 不得二次执行。已完成、失败、取消、过期任务以及普通聊天关键词均不得由补偿逻辑重新启动。",
        "发布规范：网页、私人 Bundle、测试、版本号和安装包同批更新。未做 Mac 编译、签名和真机验收时，交付说明必须明确标注。",
    ],
)

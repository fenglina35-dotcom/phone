from pathlib import Path

from docx import Document


ROOT = Path(__file__).resolve().parents[1]
DOCS = ROOT / "docs/maintenance"
MARKER = "v1062／1.0.185"


def append_section(filename, title, paragraphs):
    path = DOCS / filename
    document = Document(path)
    if any(MARKER in paragraph.text for paragraph in document.paragraphs):
        print(f"SKIP={path.name}")
        return
    document.add_heading(title, level=1)
    for paragraph in paragraphs:
        document.add_paragraph(paragraph)
    document.save(path)
    print(f"UPDATED={path.name}")


append_section(
    "AI开发项目_项目说明文档.docx",
    "v1062／1.0.185 外卖澄清动作与同任务续接修复（2026-08-25）",
    [
        "网页核心升级为 v1062，私人 iOS 升级为 1.0.185（185），原生桥继续为 25。修复真实模型输出中文全角【真实外卖｜…】时，前置检测已判定存在动作、后置执行器却只接受半角格式而丢弃动作的问题。",
        "起送价澄清由原 taskId 保存角色在当前会话列出的同店候选；用户回答“都要”或点名某一项时，只续接该任务、同一门店和原商品进度。更新前形成的旧任务仍可在同一账号和会话内受限回看最近候选。",
        "普通角色回复、朋友圈回复、后台通知、远控字幕、共同生活和外置语音配置保持原线路。最终付款、验证码和平台安全验证仍由用户本人完成。",
        "网页与私人 PhoneWeb.bundle 已同步。Windows 完成语法、自动测试和 ZIP 结构检查；Mac 编译、签名和真实 iPhone 覆盖安装仍待执行。",
    ],
)

append_section(
    "AI开发项目_Bug记录模板.docx",
    "v1062／1.0.185 外卖澄清动作与同任务续接 Bug 记录（2026-08-25）",
    [
        "现象：起送价不足后角色列出煎饺、灌汤包，用户回答“都要”，角色回复“行”，但自动化无反应，外卖服务日志没有收到新请求。",
        "根因：动作存在格式断层。检测器同时识别半角与中文全角外卖标签，因此跳过无动作兜底；真正执行器只解析半角标签。模型返回中文全角【真实外卖｜…】时，动作在最后一步被丢弃。原测试只覆盖半角动作与理想化聊天历史，未覆盖该组合。",
        "修复：执行前把中文全角真实外卖标签规范为统一半角格式；角色列出的加购候选同时保存到原 taskId，续接优先读取任务候选，并保留受限历史兼容。",
        "验证：新增全角动作规范化、候选持久化、都要续接、同 messageId 幂等、普通聊天不触发以及网页与私人 Bundle 一致性回归。",
    ],
)

append_section(
    "AI开发项目_Bug修改规范.docx",
    "v1062／1.0.185 外卖结构化动作兼容规范（2026-08-25）",
    [
        "解析规范：前置动作检测与最终执行器必须接受同一组括号和分隔符格式。任何新增标签格式都必须用端到端测试证明从检测、规范化到执行完整可达。",
        "续接规范：起送价候选必须绑定原 taskId、角色、账号和会话。用户回答“都要”只能选择该任务里由角色明确列出的候选，不能把普通聊天食物名或平台任意商品当作授权。",
        "幂等规范：同一 messageId 重复处理不得二次搜索、二次加购或新建任务；完成、取消、过期任务不得恢复。",
        "保护规范：外卖修复不得改变普通角色回复、朋友圈回复、后台通知、远控字幕、共同生活或外置语音配置；未完成 Mac 与真机验证时必须如实标注。",
    ],
)

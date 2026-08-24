from pathlib import Path

from docx import Document


ROOT = Path(__file__).resolve().parents[1]
DOCS = ROOT / "docs" / "maintenance"


def append_release(filename: str, title: str, paragraphs: list[str]) -> None:
    path = DOCS / filename
    doc = Document(path)
    if any(title in paragraph.text for paragraph in doc.paragraphs):
        return
    doc.add_page_break()
    doc.add_heading(title, level=1)
    for body in paragraphs:
        doc.add_paragraph(body)
    doc.save(path)


def append_unique_paragraph(filename: str, body: str) -> None:
    path = DOCS / filename
    doc = Document(path)
    if any(body == paragraph.text for paragraph in doc.paragraphs):
        return
    doc.add_paragraph(body)
    doc.save(path)


append_release(
    "AI开发项目_项目说明文档.docx",
    "v1054 真实外卖结构化授权与全流程稳定基线（2026-08-24）",
    [
        "网页核心升为 v1054，私人 iOS 升为 1.0.177（177），原生桥契约保持 25。网页与私人 PhoneWeb.bundle 的 delivery.js 必须逐字一致；私人包由当前网页源重新生成，保留私人主设备云备份、微信、朋友圈、普通角色回复、后台通知、远控字幕、外置语音设置及既有原生能力。",
        "每一次角色点单均由当前真实模型回合输出结构化动作并生成唯一 taskId。任务记录授权来源（用户明确要求或角色当前回合自主决定）、角色、账号、会话、消息／回合标识、创建时间、原始意图摘要、状态、已完成商品、当前澄清和修订号。旧消息、页面重绘和后台恢复只能恢复未结束任务，普通聊天关键词不能创建任务；完成、取消、过期或失败任务不能恢复。",
        "真实搜索遵循门店全局匹配→进入高吻合门店→单个商品店内搜索→单独选择杯型、冷热、糖度、加料等规格→加购→下一商品。规格不得拼进搜索词。搜索未命中可以由角色按自身口吻询问商品名，并用同一 taskId 的新 revision 续接原门店、原商品和原进度，不重新执行新任务授权。",
        "结算有两道独立门槛：用户明确商品和规格清单全部完成，以及达到平台起送价。奶茶和主食不得复制主商品凑单；奶茶只能添加不重复的同店小料，主食只能添加不重复的同店小吃。提交前再次核对真实购物车和明确排除项，并检查可用优惠券；只有已应用优惠券或明确核验为无可用券，才进入支付宝待付款。",
        "高成功率随机池如实限定为奶茶、咖啡、麦当劳、KFC、粥。水果、甜品及其他食物能力没有删除：用户明确要求时仍按真实页面尝试，角色当前回合有具体理由时也可自主选择，但不得声称与已训练类别同等稳定。麦当劳无要求套餐默认从首页选择“麦满分单人餐随心选”，在规格弹窗选择麦满分、小食和饮料；香芋派、菠萝派未指定份数时只选一份装。",
        "Windows 自动化覆盖授权幂等、终态／过期恢复拦截、当前回合结构化自主动作、普通关键词不触发、同 taskId 澄清续接、门店相似度、逐项商品和规格、起送凑单、重复商品、麦当劳／KFC／奶茶／主食／水果边界、购物车终检、优惠券状态、支付宝待付款和网页／私人一致性。Mac 编译、Apple 签名与真实 iPhone 覆盖安装仍必须单独完成；最终付款始终由本人确认。",
    ],
)

append_unique_paragraph(
    "AI开发项目_项目说明文档.docx",
    "v1054 发布复核补充：私人 bundle 生成器只保留 WKWebView 必需的 viewport、状态栏和 Service Worker 地址差异，不得改写共享 app.js 的通话身份、普通回复或其他已确认布局；资源清单显式包含真实外卖黄色袋鼠兜底图。最终生成结果要求网页与私人 app.js、delivery.js 分别内容一致。",
)

append_unique_paragraph(
    "AI开发项目_Bug修改规范.docx",
    "私人包同步不得使用过时转换回退已确认的共享功能。每次重新 stage 后必须检查网页／私人 app.js 与 delivery.js 内容一致、PhoneWeb.bundle 入口唯一，并逐项核对 manifest 中新增运行时资源；资源只在网页存在而未进入 bundle 也属于 P0 发布回归。",
)

append_unique_paragraph(
    "AI开发项目_Bug记录模板.docx",
    "发布复核追加：重新 stage 私人包后，测试发现旧 private-phone-web-transform 仍删除通话 phcallidentity 布局，manifest 也漏列 delivery-fallback-kangaroo.jpg，导致网页已修复而私人包回退。已删除过时 app／通话 CSS 转换并把兜底图加入清单；重新生成后 app.js、delivery.js 均与网页一致，兜底图存在，相关失败专项全部恢复通过。",
)

append_release(
    "AI开发项目_Bug修改规范.docx",
    "v1054：真实外卖当前回合授权、同任务续接与结算终检规范（2026-08-24）",
    [
        "授权不得从最近聊天拼接文本中反推。只有当前 actionMeta.userText 的明确用户授权，或当前真实模型返回的结构化自主点单动作，才可创建新 taskId；“饿了、没吃饭、想吃东西”等普通关键词、旧消息、页面重绘和后台恢复均不得新建任务。",
        "客户端和服务端必须共同校验 taskId、authorizationSource、roleId、accountId、sessionId、turnId、messageId、createdAt、intentSummary、status、revision 与用户约束。同一 taskId 同一 revision 的搜索和创建必须幂等；任务进入 completed、canceled、expired 或 failed 后不得恢复。",
        "商品或规格未找到不等于授权任务立即终结。客户端需要追问时，把任务置为 awaiting_clarification；用户回答后只提高原 taskId 的 revision，并从原门店、原商品和原清单进度继续。服务端搜索异常不得先把授权记录永久标为 failed，否则合法澄清会被错误拒绝。",
        "结构化商品清单是执行权威源。自然聊天原文只提供明确规格、禁选、口味和自主范围，绝不能覆盖或改写已解析商品清单；门店、商品和规格必须分层传递。相似短词可以命中带修饰词的真实单品，但不得跨类别、套餐、双杯、两份装或无关首卡。",
        "所有品类共用结算终检：明确商品与规格全部完成、购物车数量正确且无禁止项、达到起送价、优惠券检查得到 applied 或 none 四项缺一不可。平台控件异步出现时应等待并重读；不能因为金额足够就提前提交，也不能用重复主商品凑单。",
        "平台可能把同一套餐规格字幕完整重复两遍。只允许折叠完全相同且覆盖整段的重复文本，不得模糊删除相似但不同的规格；最终订单卡片、聊天结果和偏好记忆都使用规范化后的真实规格。",
        "训练状态必须如实描述。奶茶、咖啡、麦当劳、KFC、粥可作为当前高成功率随机池；水果、甜品和其他食物仍可按明确要求或角色具体判断尝试，但不得标成已训练稳定，也不得为了提高表面成功率删除能力或伪造结果。",
    ],
)

append_release(
    "AI开发项目_Bug记录模板.docx",
    "v1054 真实外卖全天训练暴露的授权、续接与结算冲突（2026-08-24）",
    [
        "现象：旧一轮“随便点”仍在最近 30 分钟消息中时，新的普通聊天可能被当成自主点单授权；搜索未命中后角色询问商品名，用户回答却可能因为服务端已把任务标为 failed 而无法续接；部分曼玲粥测试中结构化清单被自然文本干扰，出现未要求的杭州小笼包；平台套餐规格文本在订单结果里完整重复。",
        "根因一：roleRequestIntent 通过最近消息 joined 文本判断 autonomous，而不是只看当前回合。根因二：服务端 search 捕获任何异常后把授权任务设为 failed，与客户端 awaiting_clarification／revision 续接设计冲突。根因三：执行链曾让自然语言意图参与商品选择，弱化结构化清单权威性。根因四：平台确认页会在同一可见节点重复渲染规格字幕，采集器未做严格整段去重。",
        "修复：自主授权只读取当前 actionMeta.userText；角色当前回合自主决定仍要求真实模型结构化动作。搜索异常把服务端任务恢复为 authorized，允许同 taskId 提高 revision 后续接，终态任务仍严格拒绝。结构化清单逐项校验，明确商品全部完成后才凑起送价和结算。对完全相同且整段重复的规格字幕做确定性折叠。",
        "补充冲突修正：删除未调用但内容过时的七分类提示，把当前高成功率随机池从错误的六类收缩为奶茶、咖啡、麦当劳、KFC、粥；水果和甜品保留明确点单与有理由自主尝试能力，但不再虚报训练稳定度。麦当劳套餐、单份香芋派／菠萝派、奶茶加料、主食小吃、KFC 单品、优惠券与最终购物车继续使用既有专用规则。",
        "验证：网页与私人 delivery.js 哈希一致；角色学习与交接专项通过；外卖浏览服务 143/143 通过，覆盖同 taskId 失败搜索续接、幂等创建、终态／过期拒绝、普通关键词不触发、结构化自主授权、麦当劳首页套餐、单份派、奶茶小料、主食清单、购物车核对、优惠券状态和待付款边界。完成升版后继续运行项目全量测试、ZIP 文件数／哈希／单入口检查。",
        "真实页面证据：麦当劳早餐测试已到支付宝待付款，购物车仅一份“麦满分单人餐随心选”，平台总额 ¥26.90、优惠 ¥10.10、优惠券核验为无可用券；未进行付款。该结果证明流程可走通，但单次实测不能等价于 90% 成功率统计，成功率仍受门店营业、实时菜单、售罄、平台页面变化和人机验证影响。",
    ],
)

print("Updated three v1054 maintenance DOCX files")

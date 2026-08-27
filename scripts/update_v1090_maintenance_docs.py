"""Append the v1090 natural delivery commitment release to maintenance records."""

from pathlib import Path

from docx import Document
from docx.enum.text import WD_BREAK


ROOT = Path(__file__).resolve().parents[1]
DOCS = ROOT / "docs" / "maintenance"
MARKER = "v1090｜"


def append_docx(filename: str, title: str, paragraphs: list[str]) -> None:
    path = DOCS / filename
    doc = Document(path)
    existing = "\n".join(p.text for p in doc.paragraphs)
    if MARKER in existing:
        raise RuntimeError(f"{filename}: v1090 section already exists")
    page = doc.add_paragraph()
    page.add_run().add_break(WD_BREAK.PAGE)
    heading = doc.add_paragraph()
    run = heading.add_run(title)
    run.bold = True
    run.font.size = doc.styles["Normal"].font.size
    for paragraph in paragraphs:
        doc.add_paragraph(paragraph)
    doc.save(path)


def append_txt(filename: str, title: str, paragraphs: list[str]) -> None:
    path = DOCS / filename
    existing = path.read_text(encoding="utf-8")
    if MARKER in existing:
        raise RuntimeError(f"{filename}: v1090 section already exists")
    block = "\n\n" + title + "\n" + "\n".join(paragraphs) + "\n"
    path.write_text(existing.rstrip() + block, encoding="utf-8", newline="")


records = [
    (
        "AI开发项目_项目说明文档",
        "v1090｜自然外卖承诺必须落到可执行动作（2026-08-27）",
        [
            "版本：共享网页 v1090；私人 iOS 1.0.215（215）；原生桥 25。v1089 的群聊退出和关心进度，以及此前的红包兑换、订单回执、付款真值与 KFC 首页套餐规则全部保留。",
            "用户把门店和具体商品交给角色选择，例如“给我点一杯果茶，随便点一杯，不加糖就行”时，先由角色本人按完整人设决定是否接受。若角色已经明确说“行，我给你找”等承诺，执行层必须在同一授权回合补出具体真实门店、具体商品和用户硬规格，不能只留下口头答应而让后台无反应。",
            "第一次动作补判格式无效时只允许一次严格纠正；仍无唯一有效动作就显示明确失败，并确认没有连接后台或下单。普通聊天、历史陈述、拒绝和非食品安排仍不能触发外卖。",
            "浏览器界面级模拟故意让第一次补判无效，第二次严格纠正成功创建并完成一个假后端任务，选择百分茶、暴打土芭乐柠檬茶和不额外加糖；模拟调用覆盖搜索、创建订单和支付，但没有连接真实外卖后台、没有真实订单或付款。Mac 编译、签名和真实 iPhone 仍须另行验收。",
        ],
    ),
    (
        "AI开发项目_Bug修改规范",
        "v1090｜自然外卖口头承诺与执行一致性规范（2026-08-27）",
        [
            "门店和商品由角色自由选择的外卖请求不能被当成信息不足直接静默。角色可以按人设拒绝或追问；但一旦本回合已经明确答应去找、去选或去点，可见承诺就是当前决定，必须在同一授权回合生成可执行的具体门店、具体商品和明确规格。",
            "用户明确的糖度、冰度、温度等规格是硬条件，补判和严格纠正均不得漏掉或改写。用户未给规格时才允许角色根据真实可用项选择，不能把“你现在选定的”之类占位文字发送给自动化。",
            "动作补判最多一次普通补判加一次严格格式纠正。两次都失败必须 fail-closed 并给出非角色系统诊断，明确没有连接后台或下单；不得无限重试、不得制造固定角色台词，也不得把角色拒绝改成同意。",
            "测试必须同时覆盖接受、拒绝、普通聊天、用户硬规格、首次无效后的单次纠正和两次失败的显式停止；真实后台、真实下单和付款验证必须与无副作用模拟结论分开记录。",
        ],
    ),
    (
        "AI开发项目_Bug记录模板",
        "v1090｜角色答应找果茶但后台自动化无反应（2026-08-27）",
        [
            "现象：用户说“给我点一杯果茶，随便点一杯，不加糖就行”，角色回复“行，我给你找”，但外卖后台没有任务、搜索或任何可见反应。",
            "根因：门店和具体商品都交给角色选择的宽泛请求不会被显式解析器直接建任务，只依赖后续模型补判；可见的“我给你找”只被当作自然前奏，若补判模型没有输出唯一 [真实外卖|...] 动作，原流程静默结束。",
            "修复：新增受控的自由选择外卖意图和角色承诺识别。角色明确接受后，补判提示强制选择具体门店、具体商品并保留“不加糖”等硬规格；首次无效只进行一次严格纠正，仍失败则显示未启动诊断且不连接后台。",
            "验证：新增专项测试覆盖自然句、承诺、规格、严格单次纠正、拒绝与普通聊天边界；完整外卖专项 75/75 通过。浏览器界面级假后端测试中，首次补判故意无效，第二次生成有效动作并完成 search/create_order/pay_order 模拟，最终规格为不额外加糖；控制台无错误，未连接真实后台、未创建或支付真实订单。发布前还需完成根目录和外卖服务全量回归、ZIP 字节校验。",
        ],
    ),
    (
        "AI开发项目_新聊天启动说明",
        "v1090｜新聊天接手说明（2026-08-27）",
        [
            "当前候选：网页 v1090；私人 iOS 1.0.215（215）；原生桥 25。接手时先核对远端提交、安装包 SHA-256、根/私人 app.js 与 delivery.js 哈希，以及大量未跟踪资料保护状态。",
            "自然外卖自由选择请求先让角色本人决定；角色若已经明确答应去找或去点，同一回合必须生成具体门店、具体商品并保留用户明确规格。第一次动作无效只严格纠正一次，仍失败就明确停止且不连接后台。",
            "不得把普通聊天、历史陈述、拒绝或非食品安排变成外卖；不得用固定角色台词兜底，也不得以自动化测试代替真实平台、真实付款或营业门店验收。",
            "本版浏览器验证仅使用假后端模拟搜索、创建订单与支付调用；没有真实订单。Windows 完整回归和 ZIP 结构/字节校验完成后才可提交发布，Mac 编译签名和真实 iPhone 仍须单独记录。",
        ],
    ),
]


for stem, title, paragraphs in records:
    append_docx(stem + ".docx", title, paragraphs)
    append_txt(stem + ".txt", title, paragraphs)

print("Updated four maintenance DOCX/TXT pairs for v1090")

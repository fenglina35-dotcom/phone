from pathlib import Path

from docx import Document
from docx.enum.text import WD_BREAK


ROOT = Path(__file__).resolve().parents[1]
DOCS = ROOT / "docs" / "maintenance"
MARKER = "v1069／私人 1.0.193"


SECTIONS = {
    "AI开发项目_项目说明文档": (
        "v1069／私人 1.0.193 App WebContent 恢复与内存边界（2026-08-25）",
        [
            "当前网页核心保持 v1069，私人 iOS 仅升为 1.0.193（193），原生桥保持 25。本轮只处理私人 App 的 WKWebView/WebContent 稳定性；所有新增网页逻辑均以私人原生环境为入口，公开网页继续走原路径。",
            "私人 App 的存储占用明细不再逐项打开 IndexedDB 并读取每个大值，避免查看 1.68GB 图片与 WebKit 数据时把媒体复制进 JavaScript 内存。私人图片 JS 内存缓存增加 48M 字符软上限，只释放未显示、未被当前状态直接引用的内存副本，不删除 IndexedDB 图片或用户数据。",
            "WebContent 终止后不再立即 reload 当前页面。第一次终止等待 5 秒后重新载入当前安装包内唯一的 PhoneWeb.bundle/index.html；两分钟内再次终止则停止自动重载并显示失败页，避免白屏、屏保和重载循环持续发热。",
            "Windows 全量自动测试 1065／1065 通过，根 app.js 与私人 PhoneWeb.bundle/app.js 校验一致。Windows 不能证明 Mac 编译、签名或真实 iPhone 长期卡顿已修好；必须在全新 Mac 工程和原真机继续验证。",
        ],
    ),
    "AI开发项目_Bug记录模板": (
        "v1069／私人 1.0.193 混合新旧页面、白屏重载与发热（2026-08-25）",
        [
            "现象：私人 App 偶发点击延迟、快速发热、白屏后回到屏保并反复跳转；同一次异常中部分页面看起来像新版、部分像旧版。模拟器不复现，真机可能连续正常数天后再出现。",
            "已确认的代码风险一：存储占用明细原先通过 IndexedDB openCursor 读取并估算每个值，面对约 1.68GB 图片/WebKit 数据时可能把大对象复制进 WebContent；本次改为私人 App 只读取本地概要和原生容量估算。此前未打开明细也曾复现，因此不能把该扫描写成唯一根因。",
            "已确认的代码风险二：WKWebView 的 WebContent 终止后原生层最多立即 reload 两次，而网页启动又会恢复屏保状态，能够形成用户看到的白屏—屏保—白屏重载风暴。现改为一次延迟、精确本地入口恢复；第二次终止停止自动重载。",
            "长期风险：私人图片缓存原先没有上限，角色头像、壁纸和聊天图片会随页面访问留在 JavaScript 内存。本次只驱逐不再显示的内存副本，绝不清理用户存档、聊天、图片、语音或 IndexedDB。是否为多日后复现的主要原因仍须真机长时间验证。",
            "排除：私人 App 从 file URL 加载本地 PhoneWeb.bundle，Service Worker 注册会跳过，不能把混合页面简单归因于网页 Service Worker 缓存；也不能继续仅延长超时、改 CSS 或盲目增加自动 reload。",
        ],
    ),
    "AI开发项目_Bug修改规范": (
        "v1069／私人 1.0.193 WKWebView 卡顿发热修改规范（2026-08-25）",
        [
            "私人 App 出现混合页面、白屏和屏保跳转时，先区分 WebContent 终止恢复、内存增长、同步大数据扫描和页面本身未统一改版；不得把视觉差异直接写成缓存命中。",
            "WebContent 恢复必须重新打开当前安装包内确定的本地入口，且自动恢复次数和时间窗口必须有硬上限。禁止连续 reload、无界重试或用重建 WKWebView 掩盖根因。",
            "性能修复只能释放可重建的内存副本，禁止删除真实聊天、图片、语音、共同相册、朋友圈、外置语音配置、后台消息回执或原生主存档。存储诊断不得为了计算大小遍历并反序列化全部媒体值。",
            "所有共享 app.js 变更必须由私人原生环境门控，公开网页执行原路线。发布前必须校验根 app.js 与 PhoneWeb.bundle/app.js、完整自动测试、明确 iOS 构建号；Windows 通过不能代替 Mac 编译和真实 iPhone 长时间验收。",
        ],
    ),
    "AI开发项目_新聊天启动说明": (
        "v1069／私人 1.0.193 接手说明（2026-08-25）",
        [
            "当前候选：网页核心 v1069、私人 iOS 1.0.193（193）、原生桥 25。只处理私人 App 的 WebContent 恢复、存储明细大值扫描和图片内存缓存；不要修改公开网页行为。Windows 1065／1065 通过，Mac 编译、签名与本包真机验证未完成。",
            "第一步必须在 Mac 全新目录完整解压本次安装包，打开 PhoneCompanionTest.xcodeproj 编译签名并覆盖安装到原 iPhone；先确认版本页是 v1069／1.0.193／原生桥25，不能用 v1069／1.0.192 的旧安装结果代替。",
            "真机先连续点击和切换页面，再打开一次存储占用明细，观察至少 15 分钟温度、触摸、白屏和屏保跳转；随后后台放置并复测通知入聊。若 WebContent 再终止，应只延迟恢复一次，第二次停止而不是循环。",
            "必须冒烟保护普通微信、朋友圈真实回复、后台 APNs 与聊天落库、真实电量/位置/健康/屏幕时间、通话、外置语音配置和共同相册。禁止假角色回复、假设备结果、清理用户数据或以 Windows 测试冒充真机通过。",
        ],
    ),
}


def append_docx(path: Path, title: str, paragraphs: list[str]) -> None:
    doc = Document(path)
    existing = "\n".join(p.text for p in doc.paragraphs)
    if MARKER in existing:
        return
    page = doc.add_paragraph()
    page.add_run().add_break(WD_BREAK.PAGE)
    heading = doc.add_paragraph()
    run = heading.add_run(title)
    run.bold = True
    run.font.size = doc.styles["Normal"].font.size
    for paragraph in paragraphs:
        doc.add_paragraph(paragraph)
    doc.save(path)


def append_txt(path: Path, title: str, paragraphs: list[str]) -> None:
    text = path.read_text(encoding="utf-8")
    if MARKER in text:
        return
    block = "\n\n" + title + "\n" + "\n".join(paragraphs) + "\n"
    path.write_text(text.rstrip() + block, encoding="utf-8")


for stem, (title, paragraphs) in SECTIONS.items():
    append_docx(DOCS / f"{stem}.docx", title, paragraphs)
    append_txt(DOCS / f"{stem}.txt", title, paragraphs)

print("Updated four maintenance DOCX/TXT pairs for private iOS 1.0.193")

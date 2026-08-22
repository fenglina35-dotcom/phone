from pathlib import Path
from zipfile import ZipFile

from docx import Document


root = Path(__file__).resolve().parents[1] / "docs" / "maintenance"
files = sorted(root.glob("AI开发项目_*.docx"))
if len(files) != 4:
    raise RuntimeError(f"expected four maintenance DOCX files, found {len(files)}")

seen_regression = False
for path in files:
    with ZipFile(path) as archive:
        bad = archive.testzip()
        if bad:
            raise RuntimeError(f"{path.name}: corrupt member {bad}")
        if "word/document.xml" not in archive.namelist():
            raise RuntimeError(f"{path.name}: missing word/document.xml")
    doc = Document(path)
    text = "\n".join(p.text for p in doc.paragraphs)
    if "v1037｜" not in text:
        raise RuntimeError(f"{path.name}: missing v1037 section")
    seen_regression = seen_regression or "944/944" in text
    if not doc.paragraphs[-1].text.strip():
        raise RuntimeError(f"{path.name}: empty final paragraph")
    print(f"{path.name}: {len(doc.paragraphs)} paragraphs; tail={doc.paragraphs[-1].text[:72]}")

if not seen_regression:
    raise RuntimeError("maintenance set: missing full regression result")

print("v1037 maintenance DOCX structure checks passed")

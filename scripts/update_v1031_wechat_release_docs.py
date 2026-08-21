from pathlib import Path

from docx import Document
from docx.enum.text import WD_BREAK


ROOT = Path(__file__).resolve().parents[1]
DOCS = ROOT / "docs" / "maintenance"


def append_section(filename, title, paragraphs):
    path = DOCS / filename
    doc = Document(path)
    existing = "\n".join(p.text for p in doc.paragraphs)
    section_marker = title.split("（", 1)[0]
    if section_marker in existing:
        raise RuntimeError(f"{filename}: section already exists")
    page = doc.add_paragraph()
    page.add_run().add_break(WD_BREAK.PAGE)
    heading = doc.add_paragraph()
    run = heading.add_run(title)
    run.bold = True
    run.font.size = doc.styles["Normal"].font.size
    for paragraph in paragraphs:
        doc.add_paragraph(paragraph)
    doc.save(path)


append_section(
    "AI开发项目_项目说明文档.docx",
    "v1031｜微信首页视觉升级（2026-08-21）",
    [
        "共享网页从 v1029 升级为 v1031，发布标题为“微信首页视觉升级”。v1029 仍代表“真实外卖与角色钱包”；本次只为已完成的微信首页视觉修改建立新的可见网页版本。",
        "v1030 已被并行工作区的真实外卖支付回执闭环占用，本发布因此使用下一个未占用号 v1031，且不包含该并行功能。",
        "缓存标识为 v1031-wechat-home-1，Service Worker shell cache 为 north-shell-v1031。入口页、修复页、网页 shell、共享资源查询参数和私人 PhoneWeb.bundle 内嵌网页标识均同步为 1031。",
        "版本：网页 v1031；私人 iOS 仍为 1.0.150（150）；原生桥 25。Windows Node 自动复核 934/934 通过；Mac 编译、签名和真实 iPhone 验证未在本次完成。",
    ],
)

append_section(
    "AI开发项目_Bug记录模板.docx",
    "v1031｜微信更新提示未出现（2026-08-21）",
    [
        "问题：微信首页提交已进入远端 main，但用户设备已是 v1029，没有看到新版本提示；界面中的 v1029 名称仍是“真实外卖与角色钱包”。",
        "根因：首次隔离发布只更换了 hotfix 和 shell cache 标识，却保留 APP_VER、BUILD 与 __NORTH_SHELL_BUILD__ 为 1029。更新提示按数字 BUILD 比较，同为 1029 时不会向已更新设备弹出。",
        "处理：将本次微信发布独立升到 v1031，同步入口、修复页、Service Worker、静态资源查询参数、PhoneWebBundleInfo 和私人 Bundle；保留私人 iOS 1.0.150（150）和原生桥 25不变。",
        "隔离：v1030 并行支付回执工作未被合入。结果：干净 origin/main 隔离工作树 934/934 通过，版本一致性和更新提示契约测试通过。",
    ],
)

append_section(
    "AI开发项目_Bug修改规范.docx",
    "v1031｜用户可见发布必须递增 BUILD（2026-08-21）",
    [
        "需要已安装用户看到“发现新版本”时，只更改 hotfix、Service Worker cache 或查询参数不够；APP_VER、BUILD、__NORTH_SHELL_BUILD__、入口／修复页和私人内嵌网页标识必须使用同一个更大数字。",
        "多工作区并行时，发布前必须同时检查远端基线和其他工作区已占用的未发布版本号。已占用号不得重复赋予另一功能；选择下一个未占用号后，后续并行发布也必须顺延，不得降版。",
        "网页版本递增不等于私人 iOS build 必须递增。未做 Mac 编译和签名时，可仅升网页，但 PhoneWebBundleInfo 和内嵌网页内容仍须与共享网页一致，并在文档中明确写出 iOS build 未变。",
    ],
)

append_section(
    "AI开发项目_新聊天启动说明.docx",
    "v1031｜微信首页可见更新交接（2026-08-21）",
    [
        "当前网页发布基线为 v1031，私人 iOS 仍为 1.0.150（150），原生桥 25。v1031 的用户可见名称是“微信首页视觉升级”；v1029 的“真实外卖与角色钱包”仍作为其前置基线保留。",
        "本次解决了已在 v1029 的设备无法收到微信视觉更新提示的问题。缓存标识 v1031-wechat-home-1，shell cache north-shell-v1031，并已同步入口页、修复页和私人 PhoneWeb.bundle。",
        "另一工作区的 v1030 真实外卖支付回执闭环不在本发布中。该工作后续上线时必须合入 v1031 基线并改用更高网页版本，不能将线上从 v1031 降回 v1030。",
        "Windows Node 自动复核 934/934 通过。Mac 编译、签名和真实 iPhone 验证仍未完成。",
    ],
)

print("Updated four maintenance DOCX files for v1031 WeChat release")

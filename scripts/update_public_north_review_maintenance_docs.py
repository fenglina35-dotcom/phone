from pathlib import Path

from docx import Document


ROOT = Path(__file__).resolve().parents[1]
DOCS = ROOT / "docs" / "maintenance"


def append_section(filename: str, title: str, paragraphs: list[str]) -> None:
    path = DOCS / filename
    document = Document(path)
    if any(paragraph.text.strip() == title for paragraph in document.paragraphs):
        return
    document.add_heading(title, level=1)
    for text in paragraphs:
        document.add_paragraph(text)
    document.save(path)


append_section(
    "AI开发项目_项目说明文档.docx",
    "2026-08-21｜公开 North App Review 5.6 透明审核修复候选",
    [
        "本轮处理公开 North 1.0 因 Guideline 5.6 被拒的问题。Apple 的拒绝信息认为审核过程中可能存在被隐藏的功能，但源码核对未发现按 App Review、TestFlight、地区、日期、设备类型或模拟器隐藏功能的判断。真实风险是核心“角色远程管理”虽然在 App 内以“真实同步”出现，但审核员没有可用控制端、永久演示账号、预创建角色和完整测试步骤，因而无法独立完成配对与远程命令闭环。",
        "已从被拒构建对应的完整 PhoneCompanionTest 源码建立隔离目录 native/public-north-review。公开工程继续只包含本机管理、Family Controls／Managed Settings／Device Activity、可选健康与位置、短时配对、角色命令和设备回执，不嵌入私人“小手机”网页。主 App 与四个扩展保持原 Bundle ID、App Group 和 entitlement 边界，所有 Target 的构建号明确提升为 5。",
        "App 界面把“管理”明确改为“本机管理”，把“真实同步”明确改为“角色远程管理”。远程页公开说明：角色控制端可以发送查看、锁定、解锁和每日限额命令；只执行用户在 Apple 系统选择器中明确选择的 App；服务器排队不等于设备成功，必须等待设备回执。App 内提供公开角色控制台链接，并说明审核账号与测试角色由 Review Notes 提供。",
        "新增公开角色控制台 north-role-controller.html／js 和迁移 202608210001_north_review_portal.sql。审核员使用永久 Email／Password 演示账号登录，测试角色和目标已由开发者预配置；无需注册或创建角色。每次配对仍生成十分钟有效的一次性 8 位码。审核 RPC 只允许 authenticated 用户访问 auth.uid() 绑定的目标；网页不返回位置和健康快照，也不接触 owner secret、device secret 或 service-role key。",
        "为兼容曾注册旧小手机 Service Worker 的浏览器，正式审核 URL 使用从 v876 起已排除应用外壳缓存的 north-support.html?role-controller=1，再以 no-store 读取同源控制台。支持页、隐私页和当前 Service Worker 同时公开角色控制能力和数据边界；不存在暗门、固定万能码或审核设备识别。",
        "验证：角色控制台 JavaScript 语法通过；审核页面／RPC／Swift 透明性专项 7／7 通过；桌面和 390×844 手机页面实测无控制台错误；旧 Service Worker 环境可由稳定支持页入口正确打开控制台。本工作区完整 Node 回归 910 项中 908 通过，2 个失败来自当前已有 v1018 聊天尾部持久化与 storage-overflow 工作，不属于本轮文件，未越权修改。Windows 无法执行 Xcode 编译、签名、Supabase 实际迁移、创建真实审核账号或 iPhone 真机回执，以上仍待 Mac／云端完成。",
    ],
)

append_section(
    "AI开发项目_Bug记录模板.docx",
    "2026-08-21｜公开 North 5.6 拒绝｜角色远程管理无法由审核员独立验证",
    [
        "现象：App Store Connect 拒绝 North 1.0，理由为 Guideline 5.6，称 App 出现通常与欺诈相关的异常行为，并认为有功能在审核过程中被故意隐藏。审核消息没有指出具体页面、步骤或截图；开发者回复询问后未收到进一步说明。",
        "事实核对：被拒构建只包含公开 PhoneCompanionTest 伴生系统，没有完整私人“小手机”网页。界面始终显示“管理”和“真实同步”两个标签；源码中未发现 App Review、TestFlight、地区、模拟器、日期或设备型号门禁。真实同步要求 yb_ 目标和 8 位短时配对码，配对后可接收角色查看、锁定、解锁和限额命令。",
        "根因判断：审核员只能看到需要外部小手机 ID 和配对码的入口，却没有可登录控制端、永久演示账号、预创建角色和可重复生成的新配对码，因此核心远程角色能力不可审核。原 Review Notes 只解释了本机 Screen Time 能力和可选网页配对，没有给出完整可执行访问，容易被理解为申报功能与审核期间可见行为不一致。",
        "修复：建立公开 North 审核工程副本，公开改名并解释角色远程管理；新增登录保护的角色控制台、auth.uid() 绑定的预配置测试角色、十分钟一次性配对码、刷新／锁定／解锁／限额命令和设备回执列表；新增完整英文 Review Notes 模板和部署清单。支持页与隐私页同步披露角色控制端账户及命令边界。",
        "旧缓存失败方案：只把新控制台文件加入 Service Worker 白名单仍不能保证第一次访问。已经被旧 Worker 控制的浏览器会在新 Worker 更新前先把新地址换成小手机主页。本轮改用旧版本已经长期放行的 north-support.html 稳定路径作为正式控制台入口，再 no-store 加载同源控制台，旧缓存实际复测通过。",
        "安全边界：不删除角色远程管理，不上传私人 App，不加入审核专用暗门或万能配对码，不把账号密码写进 Git。控制台只向当前登录审核账号返回屏幕使用时间与所选 App 状态，不返回位置或健康快照；命令只有设备回执后才算完成。",
        "验证结果：专项 7／7、JavaScript 语法、git diff --check、桌面与手机页面、旧 Service Worker 入口均通过。完整 Node 为 908／910，两个失败属于并行 v1018 持久化改动。Supabase 部署、永久审核账号、Mac 五 Target 编译、真实 iPhone 授权／配对／锁定／解锁／限额和 App Store 重新提交尚未完成，不能写成已上架或已通过审核。",
    ],
)

append_section(
    "AI开发项目_Bug修改规范.docx",
    "公开 North 审核访问与核心远程功能透明规范（2026-08-21）",
    [
        "公开 App 的核心能力依赖外部控制端、账号、角色或配对时，Review Notes 必须提供永久可用的演示登录、无需审核员自行注册的预创建数据、可重复生成的新短时配对码和从授权到设备回执的逐步操作。只解释功能存在但不给可执行访问，视为审核链路未完成。",
        "不得通过删除角色远程控制来换取通过，也不得上传私人完整小手机替代公开 North。公开 North、私人小手机和网页版的产品边界、Bundle ID、控制器租约、云项目及用户数据必须保持隔离；公开 North 与私人小手机不得同时控制同一台本人设备。",
        "审核友好不能通过审核设备识别、TestFlight 判断、地区／日期开关、隐藏手势、固定万能码或仅审核时出现的界面实现。产品正常用户与审核员应看到同一项“角色远程管理”能力；差别只允许是审核员使用单独的受保护演示账号和预配置测试数据。",
        "审核登录不得把邮箱、密码、owner secret、device secret、service-role key、JWT 或永久配对码写入前端、文档仓库或 App 二进制。浏览器只保留当前内存会话；服务端按 auth.uid() 限定一个预配置目标；配对码继续单次、短时失效。",
        "角色远程命令必须区分服务器排队与设备执行。控制台、App 和 Review Notes 都要明确：只有设备应用 Managed Settings 后回传终态和新快照，才能显示完成；pending、queued 或 APNs 接受均不能冒充锁定／解锁成功。",
        "在 PWA 根作用域新增审核入口时，必须测试已被旧 Service Worker 控制的首次导航。仅修改新 Worker 白名单不足以修复旧 Worker 的第一次拦截；优先复用历史版本已经放行的稳定公共路径，或提供不依赖旧 Worker 更新的网络入口，并用旧缓存环境实际验证。",
    ],
)

append_section(
    "AI开发项目_新聊天启动说明.docx",
    "新聊天接手状态｜公开 North 5.6 透明审核修复候选（2026-08-21）",
    [
        "公开 North 的被拒源码已隔离导入 native/public-north-review，候选版本仍为 1.0，所有 Target 构建号为 5。它只包含本机 Screen Time 管理和公开的角色远程管理，不包含私人小手机网页；不得用 native/private-small-phone 或当前私人安装包覆盖。",
        "审核方案已落地为永久演示登录＋预创建 North Review Role＋每次十分钟一次性配对码。正式控制台 URL 为 https://fenglina35-dotcom.github.io/phone/north-support.html?role-controller=1；账号和密码只填 App Store Connect Review Notes，仓库模板必须保持占位符。",
        "服务端新增迁移 202608210001_north_review_portal.sql，但当前尚未部署到公开 North 使用的旧业务项目，也尚未创建实际审核 Auth 用户和角色映射。接手后不得先提交 App 审核；应按 docs/north-app-review-deployment.md 完成云端、账号和真机闭环。",
        "Mac 下一步：打开 native/public-north-review/PhoneCompanionTest/PhoneCompanionTest.xcodeproj，核对签名和全部五个 Target，编译安装构建 5；在干净 iPhone 完成授权、选择 App、控制台配对、上传、刷新、锁定、解锁和 15 分钟限额，并逐项等待设备回执。Windows 的专项通过不能代替这些结果。",
        "不可破坏：私人 App 与公开 North 隔离、短时一次性配对、角色／设备秘密分离、服务器排队与设备完成分离、Support／Privacy／Controller 不被应用 Shell 替换、位置与健康数据不返回审核控制台，以及现有 v1018 并行工作。完整回归当前 908／910，两个既有失败不属于公开 North 修复。",
    ],
)

print("Updated four maintenance DOCX files for public North review")


append_section(
    "AI开发项目_项目说明文档.docx",
    "2026-08-21｜公开 North 审核云端已部署",
    [
        "公开 North 审核控制端已部署到旧业务项目 lkhlyfpssmrjkkzhuzag。选择依据不是项目名称，而是线上结构核验：该项目仍具备 phone_companion_links、phone_companion_commands、绑定与短时配对 RPC，且当前公开 North 构建和角色控制台均连接此项目；邀请码项目 lovbzibismsjqvjujilz 不具备伴生表和 RPC，未被混用；私人伴生云 qvuahlqimcfgeoetosnl 未被修改。",
        "迁移 202608210001_north_review_portal.sql 已执行并登记到远端 migration history。远端核验发现 hosted Supabase 默认权限曾显式给 anon 角色新函数 EXECUTE，单独从 PUBLIC revoke 不足；迁移已修正为同时从 PUBLIC 与 anon 撤销三个审核 RPC 的执行权。复验结果为 anon 三项全部 false、authenticated 三项全部 true，审核映射表 RLS 已开启。",
        "永久审核 Auth 用户、North Review Role 和随机 yb_ 测试目标已经建立。审核密码只保存在 Windows 凭据管理器 North App Review:lkhlyfpssmrjkkzhuzag，不写入 Git、网页、SQL、维护文档或聊天交接。账号登录及 phone_companion_review_session 已在线验证成功；当前 linked=false 是尚未与候选 iPhone 构建配对的正确状态。",
        "GitHub Pages 的支持页兼容入口、直接控制台、控制台脚本和隐私页均在线返回 HTTP 200。仍未完成：Mac 五 Target 编译签名、干净 iPhone 授权与选 App、生成十分钟配对码、刷新／锁定／解锁／15 分钟限额设备回执、完整录屏和 App Store Connect Review Notes／重新提交。",
    ],
)

append_section(
    "AI开发项目_Bug记录模板.docx",
    "2026-08-21｜公开 North 审核 RPC 默认匿名执行权限",
    [
        "现象：202608210001 迁移首次部署后，审核表、三个 RPC 与 RLS 均存在，但远端 has_function_privilege 核验显示 anon 和 authenticated 都能执行三个 security definer 审核函数。",
        "根因：该 hosted Supabase 项目的函数默认权限显式包含 anon=X。原迁移只执行 REVOKE ... FROM PUBLIC；撤销 PUBLIC 不会自动移除已经直接授予 anon 的权限，因此仅凭源码中的 PUBLIC revoke 不能证明匿名访问已关闭。",
        "修复：三个函数统一改为 REVOKE ALL ... FROM PUBLIC, anon，再只向 authenticated 授予 EXECUTE；重新执行同一幂等迁移。专项测试增加对三条显式 anon revoke 的断言。",
        "线上验证：review_accounts 表存在且 RLS=true；三个 RPC 均存在；anon_session、anon_pairing、anon_command 全部 false；authenticated 对应三项全部 true。账号密码登录与绑定到 auth.uid() 的 review_session 返回 200。",
    ],
)

append_section(
    "AI开发项目_Bug修改规范.docx",
    "Supabase Security Definer RPC 权限验收规范（2026-08-21）",
    [
        "新建 security definer RPC 时不得假设 REVOKE FROM PUBLIC 等价于禁止 anon。Hosted 项目可能通过 default privileges 直接给 anon EXECUTE；迁移应显式 REVOKE ALL ON FUNCTION ... FROM PUBLIC, anon，再向所需角色精确 GRANT。",
        "部署完成必须在真实远端使用 has_function_privilege 分别核对 anon 与 authenticated，而不能只看 SQL 文本、RLS 状态或函数是否存在。审核账号绑定类 RPC 的合格条件是 anon=false、authenticated=true，并且所有数据选择继续以 auth.uid() 约束。",
    ],
)

append_section(
    "AI开发项目_新聊天启动说明.docx",
    "新聊天接手状态｜公开 North 审核云端完成（2026-08-21）",
    [
        "202608210001_north_review_portal.sql 已部署并登记到 lkhlyfpssmrjkkzhuzag；不要再部署到邀请码项目 lovbzibismsjqvjujilz，也不要把公开审核数据迁入私人伴生云 qvuahlqimcfgeoetosnl。旧 North 的配对表和 RPC 已在线核验可用。",
        "远端曾发现 Supabase 默认权限直接给 anon EXECUTE，现已修正为三个审核 RPC 明确撤销 PUBLIC 与 anon，只允许 authenticated。线上 has_function_privilege 复验全部符合预期。",
        "永久审核 Auth 用户、预建 North Review Role 和随机目标已创建，密码只在 Windows 凭据管理器 North App Review:lkhlyfpssmrjkkzhuzag 中。不要把邮箱或密码提交到 Git；仅在最终 App Store Connect Review Notes 中填写。",
        "下一步必须在 Mac 编译构建 5，并用干净 iPhone 完成授权、选 App、配对、真实数据上传、刷新、锁定、解锁和 15 分钟限额回执。当前云端 session 显示 linked=false，只有真机配对完成后才能变为 true。",
    ],
)

print("Updated four maintenance DOCX files for public North cloud deployment")

from pathlib import Path

from docx import Document


ROOT = Path(__file__).resolve().parents[1]
DOCS = ROOT / "docs" / "maintenance"


def append_section(path: Path, heading: str, paragraphs: list[str]) -> None:
    doc = Document(path)
    if any(p.text.strip() == heading for p in doc.paragraphs):
        raise RuntimeError(f"section already exists: {heading}")
    doc.add_heading(heading, level=1)
    for paragraph in paragraphs:
        doc.add_paragraph(paragraph)
    doc.save(path)


append_section(
    DOCS / "AI开发项目_项目说明文档.docx",
    "2026 09 06 私人共同生活键盘坐标修复候选",
    [
        "发布范围：私人内置网页 v1189／私人 iOS 1.0.315 (315)，原生桥 35；公开网页仍为 v1184，公开网页源码和微信输入区不改动。",
        "真机纠错：v1188 删除触摸开始阶段的强制滚动后，实际 iPhone 的共同生活光标仍偏，首次打开仍跳，收起仍由键盘和输入栏分两拍完成，因此 v1188 不能继续记为有效修复。",
        "重新对照：当前正常的微信输入区和网页版共同生活都让包含输入区的 phone 外壳从 fixed 切到同尺寸 absolute；私人 Bundle 从 v1187 起只把 offinput 从该规则删除。fixed 祖先正是代码注释中记录的 iOS WebKit 光标坐标脱节条件。",
        "v1189 候选修复：只在私人 Bundle 恢复 offinput 与 chat-inputbar 共用的 absolute 坐标基准；保留单一 WKWebView 键盘宿主，不增加键盘通知、visualViewport 监听、强制滚动或重复 focus。旁白触摸另加延迟合成点击拦截，避免误开后方消息编辑删除。",
        "验证边界：Windows 自动测试只验证代码路径、版本、公开网页零改动和安装包内容。真实光标及动画仍须 Mac/Xcode 编译签名后在实际 iPhone 连续验收。",
    ],
)

append_section(
    DOCS / "AI开发项目_Bug记录模板.docx",
    "v1189 私人 iOS315 共同生活光标与键盘同步记录",
    [
        "真机现象：v1188 中共同生活输入框点字位置与光标位置不一致；首次打开页面会咔哒上跳，收起键盘时键盘先下去，输入栏随后回落；点击旁白还可能误开附近消息的编辑删除。微信同机正常。",
        "已排除：v1188 已删除 offComposerPinLatest，现象仍完整存在，所以触摸时滚动改写不是充分根因；当前原生 SmallPhonePrivateRootView 和 LocalPhoneWebView 与早期基线没有新的共同生活专属键盘接管。",
        "定位结果：私人 Bundle 的 phone 只在 chat-inputbar 出现时切换 absolute，共同生活 offinput 仍留在 fixed 祖先中；公开网页和早期修复规则同时覆盖 chat-inputbar 与 offinput。该差异与同机微信正常、共同生活光标及位移异常一致。",
        "v1189 候选处理：恢复私人 offinput 的 absolute 祖先坐标基准，不改微信 DOM、CSS、脚本或公开网页；触摸旁白后的 1600 毫秒内拦截可能被 iOS 重定向到 offmsg 或 offnar 的合成 click。",
        "真机验收：输入框左中右三点定位；首次打开；连续开关五轮；旁白与普通模式互切五轮；紧贴消息时点旁白十次；最后单独复测微信。任何一项失败都不得写成完成。",
    ],
)

append_section(
    DOCS / "AI开发项目_Bug修改规范.docx",
    "新增规则 同机正常输入区应先做坐标系差分",
    [
        "同一 WKWebView 内若微信输入正常而共同生活异常，优先比较从 html、body、phone、page 到 textarea 的 fixed、absolute、transform、overflow 和 flex 链；在找出差分前不得继续叠加原生键盘通知或动画延时。",
        "修复光标坐标时，应让异常输入区复用正常输入区的几何原则，不复制业务 DOM 或视觉样式。必须证明公开网页、微信消息逻辑和其他页面没有被改写。",
        "键盘打开与收起出现两拍位移时，视为存在两个几何所有者。只保留 WebKit 与一个稳定祖先坐标基准；禁止同时使用 visualViewport、键盘 frame、scrollTop、contentOffset 或重复 focus 做补偿。",
        "输入栏按钮在键盘重排时必须阻止事件穿透。除 touchstart 或 pointerdown 去重外，还要覆盖 iOS 延迟 synthetic click，并仅在刚发生输入栏触摸时拦截后方消息，不能影响正常消息编辑。",
    ],
)

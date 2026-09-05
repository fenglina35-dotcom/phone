from pathlib import Path

from docx import Document


ROOT = Path(__file__).resolve().parents[1]
DOCS = ROOT / "docs" / "maintenance"


def append_section(path: Path, heading: str, paragraphs: list[str]) -> None:
    doc = Document(path)
    if any(p.text.strip() == heading for p in doc.paragraphs):
        raise RuntimeError(f"section already exists: {heading}")
    doc.add_heading(heading, level=1)
    for paragraph in paragraphs:
        doc.add_paragraph(paragraph)
    doc.save(path)


append_section(
    DOCS / "AI开发项目_项目说明文档.docx",
    "2026 09 06 私人键盘基线恢复候选",
    [
        "发布范围：仅私人 iOS 壳层及其内置 PhoneWeb.bundle，候选版本 v1187／1.0.313 (313)，内置网页业务基线仍为 v1184。公开网页源码、网页发布文件、微信业务逻辑和共同生活业务逻辑均不改动。",
        "版本复核结论：v1179 使用稳定的全高 WKWebView、SwiftUI 忽略键盘底部安全区、私人 viewport 不声明 interactive-widget，并且只有微信聊天页退出 fixed 坐标系。v1181 将线下输入栏并入同一定位规则；v1182 与 v1184 改由 viewport 和 SwiftUI 调整高度；v1185 增加 UIKeyboardLayoutGuide；v1186 又在聚焦期间开关外层滚动。后四种布局干预相互叠加，和真机上的微信回归、首次弹跳、收起分段及旧消息停留相符。",
        "候选修复：恢复 v1179 的单一键盘布局权威，但不整版回退。保留后来有效的 16 像素输入字体、触摸去重和旁白切换不重复 focus；删除线下输入栏的 absolute 祖先规则、私人 interactive-widget、原生键盘导轨、键盘通知和滚动开关。用户点线下输入框时，在键盘出现前先把页面内部的消息区锚定到最后一条，避免从旧消息位置开始上推。",
        "验证边界：Windows 可验证源码作用域、版本矩阵、公开网页零改动、测试和安装包内容；真实动画必须在 Mac/Xcode 编译签名后于实际 iPhone 验证。验收必须同时观察微信与共同生活，且确认聚焦时最后一条消息仍在输入栏上方，不能只检查黑条是否消失。",
    ],
)

append_section(
    DOCS / "AI开发项目_Bug记录模板.docx",
    "v1187 私人 iOS313 键盘回归记录",
    [
        "最新真机证据：v1186 只消除了部分黑色空区，微信输入框仍未恢复；共同生活首次点击仍上下弹跳，收起时键盘先落下、输入栏随后才落下，且点击输入框后显示的是前面几条消息，不是从最后一条消息整体上推。",
        "逐版结果：v1179 的微信键盘正常；v1180 改善线下字体和触摸；v1181 把 .offinput 加入微信使用的 absolute 定位规则；v1182 首次加入 interactive-widget 并取消 SwiftUI 键盘忽略；v1183 又恢复键盘忽略且旁白不再二次 focus；v1184 再次切回 viewport 和 SwiftUI 调整；v1185 用全局 UIKeyboardLayoutGuide 同时影响微信；v1186 撤回导轨但以键盘通知和外层 scrollView 开关继续干预线下。",
        "失败方案判定：v1185 的全局键盘导轨和 v1186 的聚焦滚动开关均未通过真实 iPhone 验收，不再继续微调。仅恢复某一个属性不足以修复，因为 viewport、SwiftUI、fixed 或 absolute 坐标系、WebKit 自动聚焦滚动和消息区滚动位置仍可能同时参与。",
        "v1187 处理：让私人 WKWebView 回到 v1179 的单一布局路径；微信恢复 v1179 的 CSS 和宿主契约；线下不再套用微信的 absolute 规则。只新增一个私人网页内部动作，在 off_in 的触摸开始时同步把 offbg 锚定到末尾，不监听键盘动画、不修改 WKScrollView、不触发额外 focus。",
        "当前状态：代码与自动化测试通过后可生成 Mac 待编译包；Windows 无法证明真实 iPhone 动画已经修好。真机必须覆盖首次打开、连续开关、滑到旧消息后点击输入框、旁白互切，以及微信进入、输入、收起和再次打开。",
    ],
)

append_section(
    DOCS / "AI开发项目_Bug修改规范.docx",
    "新增规则 键盘基线对照与最后消息验收",
    [
        "修复键盘问题前必须建立逐版本矩阵，至少记录宿主安全区、viewport、fixed 或 absolute 定位、原生键盘约束、焦点调用和滚动控制。不得只拿一个旧版本整版覆盖，也不得在失败方案上继续叠加补丁。",
        "选择最终方案时，应从已知正常版本恢复一条完整布局链，再逐项保留后来与布局无关且已验证有效的改动。每个保留项必须说明它解决的现象及为什么不会重新引入共享页面回归。",
        "输入框验收不只看位置。聚焦前消息区必须位于最后一条；键盘出现时应由最后一条消息和输入栏一起稳定上移；收起时键盘、输入栏和内容应在同一系统动画内复位；首次与后续操作都不得出现额外弹跳、黑屏或停在旧消息。",
        "私人键盘修改必须反向检查微信、共同生活、线下约会、旁白切换和其他输入页。公开网页不在修复范围时必须保持字节零改动；私人包内也不得用共享原生容器对所有页面追加键盘通知、滚动开关、偏移重写或第二套动画时间线。",
    ],
)

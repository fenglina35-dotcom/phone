from pathlib import Path

from docx import Document
from docx.enum.text import WD_BREAK


ROOT = Path(__file__).resolve().parents[1]
DOCS = ROOT / "docs" / "maintenance"
MARKER = "v1070／私人 1.0.194"


SECTIONS = {
    "AI开发项目_项目说明文档": (
        "v1070／私人 1.0.194 版本纠正与稳定边界（2026-08-25）",
        [
            "当前网页核心为 v1070，私人 iOS 为 1.0.194（194），原生桥保持 25。本次只纠正上一候选包沿用 v1069／1.0.193 所造成的识别混淆，不增加新的功能逻辑。",
            "私人 App 的存储明细安全汇总、图片内存软上限、WebContent 一次延迟恢复和两分钟内第二次终止停止重载均保持原实现；普通微信、朋友圈真实回复、后台通知入聊、真实设备数据、通话、外置语音、共同相册和远控保护线不得改写。",
            "版本壳层、私人 Bundle、缓存标识和 iOS 构建号必须成套一致。Windows 自动测试、ZIP 一致性校验、Mac 编译签名和真实 iPhone 验收仍是彼此独立的结论。",
        ],
    ),
    "AI开发项目_Bug记录模板": (
        "v1070／私人 1.0.194 候选包版本号纠正（2026-08-25）",
        [
            "现象：上一份已经包含私人 App 稳定修正的新包仍显示网页核心 v1069，只通过 iOS 1.0.193 区分，用户无法直观确认安装的是不是新候选包。",
            "处理：保留 v1069／1.0.193 的源码提交、旧 ZIP 和安装说明作为恢复资料，另行生成 v1070／1.0.194 独立包；不覆盖或删除旧交付物。",
            "边界：本次不得借版本升级继续修改卡顿逻辑或核心功能。若真机仍复现，必须按 v1070／1.0.194 的日志和页面证据另行诊断。",
        ],
    ),
    "AI开发项目_Bug修改规范": (
        "v1070／私人 1.0.194 发布编号规范（2026-08-25）",
        [
            "已经公开给用户识别或已经生成安装包的候选版本，不得再次沿用同一网页版本和同一私人 iOS 构建号。纠正包必须同时拥有新的网页核心号、iOS marketing version、build number、缓存号、安装说明和独立交付目录。",
            "纯版本纠正只能做机械标识更新和回归测试，不得混入业务逻辑修改。旧安装包、Git 提交和维护记录必须保留，禁止用同名新包静默覆盖历史恢复资料。",
        ],
    ),
    "AI开发项目_新聊天启动说明": (
        "v1070／私人 1.0.194 接手说明（2026-08-25）",
        [
            "当前候选必须显示网页核心 v1070、私人 iOS 1.0.194（194）、原生桥 25。v1069／1.0.193 是保留的上一候选，不得把两者写成同一个安装结果。",
            "本次仅做版本纠正；私人 App 的 WebContent 恢复与内存边界沿用上一候选实现。接手后先核对包内壳层、Bundle、iOS 构建和 SHA，再在 Mac 全新目录编译并覆盖安装到真实 iPhone。",
            "禁止为了版本升级触碰普通微信、朋友圈、后台通知、真实数据、通话、外置语音、共同相册和远控。Windows 测试通过不能写成 Mac 编译或真机已经修好。",
        ],
    ),
}


def append_docx(path: Path, title: str, paragraphs: list[str]) -> None:
    doc = Document(path)
    existing = "\n".join(p.text for p in doc.paragraphs)
    if MARKER in existing:
        return
    page = doc.add_paragraph()
    page.add_run().add_break(WD_BREAK.PAGE)
    heading = doc.add_paragraph()
    run = heading.add_run(title)
    run.bold = True
    run.font.size = doc.styles["Normal"].font.size
    for paragraph in paragraphs:
        doc.add_paragraph(paragraph)
    doc.save(path)


def append_txt(path: Path, title: str, paragraphs: list[str]) -> None:
    text = path.read_text(encoding="utf-8")
    if MARKER in text:
        return
    block = "\n\n" + title + "\n" + "\n".join(paragraphs) + "\n"
    path.write_text(text.rstrip() + block, encoding="utf-8")


for stem, (title, paragraphs) in SECTIONS.items():
    append_docx(DOCS / f"{stem}.docx", title, paragraphs)
    append_txt(DOCS / f"{stem}.txt", title, paragraphs)

print("Updated four maintenance DOCX/TXT pairs for v1070 / private iOS 1.0.194")

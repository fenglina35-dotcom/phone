"""Append the v1080 real-delivery role acknowledgement release to maintenance records."""

from pathlib import Path

from docx import Document
from docx.enum.text import WD_BREAK


ROOT = Path(__file__).resolve().parents[1]
DOCS = ROOT / "docs" / "maintenance"
MARKER = "v1080｜"


def append_docx(filename: str, title: str, paragraphs: list[str]) -> None:
    path = DOCS / filename
    doc = Document(path)
    existing = "\n".join(p.text for p in doc.paragraphs)
    if MARKER in existing:
        raise RuntimeError(f"{filename}: v1080 section already exists")
    page = doc.add_paragraph()
    page.add_run().add_break(WD_BREAK.PAGE)
    heading = doc.add_paragraph()
    run = heading.add_run(title)
    run.bold = True
    run.font.size = doc.styles["Normal"].font.size
    for paragraph in paragraphs:
        doc.add_paragraph(paragraph)
    doc.save(path)


def append_txt(filename: str, title: str, paragraphs: list[str]) -> None:
    path = DOCS / filename
    existing = path.read_text(encoding="utf-8")
    if MARKER in existing:
        raise RuntimeError(f"{filename}: v1080 section already exists")
    block = "\n\n" + title + "\n" + "\n".join(paragraphs) + "\n"
    path.write_text(existing.rstrip() + block, encoding="utf-8", newline="")


records = [
    (
        "AI开发项目_项目说明文档",
        "v1080｜真实外卖角色回执与付款事实边界（2026-08-27）",
        [
            "版本：共享网页 v1080；私人 iOS 1.0.205（205）；原生桥 25。v1079 的图文照片卡、吃货豆红包兑换和此前修复继续保留。",
            "真实订单提交后，订单卡片只展示真实明细；订单成功事实通过可靠功能事件队列交给角色，角色必须按本人完整人设使用第一人称“我”自然说明是自己点的、已经点好或下过单，不能只发卡片后沉默。",
            "该回执在角色忙碌、线下或通话等状态下可以延后但不能作为普通后台主动消息丢弃。第一次真实模型没有可见回复时，只沿用既有规则做一次真实模型重试，不增加固定角色兜底台词。",
            "付款状态必须以平台真实状态为准：pending_payment 只代表订单已提交，角色只能说已经点好或下单；只有 paid 或后续配送状态才可说已经付好。角色不得催用户“快去付款”、命令付款，也不得提卡片、二维码、待付款等界面词。",
        ],
    ),
    (
        "AI开发项目_Bug修改规范",
        "v1080｜订单卡片后角色回执与付款真值规范（2026-08-27）",
        [
            "真实订单卡片和角色自然回复是两个职责：卡片负责展示真实订单字段，角色回复负责以本人第一人称承认自己完成点单。不得用“卡片已展示结果”诱导模型沉默。",
            "订单提交事实必须走不会被普通 proactive 忙碌门禁丢弃的可靠功能事件队列；允许延后，不允许伪造固定台词。真实模型首次空输出时最多按现有机制重试一次。",
            "automatic:false 或 pending_payment 不能改写成 paid。只有平台 status 达到 paid 或后续配送状态，才允许角色声称付款完成；否则只能说已点好或已下单。",
            "角色回执严禁发出“自己去付款”“快去付款”等命令，不得暴露卡片、二维码、付款入口、收银台或待付款等内部界面词。",
        ],
    ),
    (
        "AI开发项目_Bug记录模板",
        "v1080｜真实订单只发卡片且催付款（2026-08-27）",
        [
            "现象：角色完成真实外卖后只显示订单卡片，没有用自己的口吻说已经点好；部分话术还可能催用户“快去付款”。角色忙碌、线下或通话时，卡片后的普通回复可能被丢弃。",
            "根因：卡片先作为 assistant 消息写入后，resultReply 后续事实被 scheduleReply 判断为普通 proactive；terminal prompt 又写着卡片已负责展示真实结果，诱导模型保持沉默。",
            "修复：把订单提交事实包装为 featureEventNote 进入可靠功能事件队列；提示要求第一人称自然承认本人点单，禁止付款命令和界面词。依据真实 status 单独提供付款事实，pending_payment 绝不视为已付款。",
            "验证：根目录全量、外卖服务以及 chat-card、role-handoff、state-machine、mode、payment-return-sync 专项均需在发布前通过；根/私人 delivery.js 必须逐字节一致。没有为本修复创建真实订单。",
        ],
    ),
    (
        "AI开发项目_新聊天启动说明",
        "v1080｜新聊天接手说明（2026-08-27）",
        [
            "当前候选：网页 v1080；私人 iOS 1.0.205（205）；原生桥 25。先核对远端提交、ZIP SHA-256、根/私人 delivery.js 哈希和大量未跟踪资料保护状态。",
            "真实外卖成功后必须同时具备真实订单卡片和角色本人第一人称自然回执；回执可因忙碌延后，但不能作为普通主动消息丢弃，也不能用固定系统台词冒充角色。",
            "pending_payment 只能说已经点好或下单，不能说已经付款。角色不能命令或催促用户付款，只有平台 paid 或后续配送状态才允许说明已付。",
            "Windows 自动测试和 ZIP 结构校验不能写成 Mac 编译、签名或真实 iPhone 已通过；普通角色回复、朋友圈、后台通知、远控字幕和外置语音配置继续作为发布保护线。",
        ],
    ),
]

for stem, title, paragraphs in records:
    append_docx(stem + ".docx", title, paragraphs)
    append_txt(stem + ".txt", title, paragraphs)

print("Updated four maintenance DOCX/TXT pairs for v1080")

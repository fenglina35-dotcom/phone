from pathlib import Path

from docx import Document
from docx.enum.text import WD_BREAK


ROOT = Path(__file__).resolve().parents[1]
DOCS = ROOT / "docs" / "maintenance"


def append_section(filename, title, paragraphs):
    path = DOCS / filename
    doc = Document(path)
    existing = "\n".join(p.text for p in doc.paragraphs)
    section_marker = title.split("（", 1)[0]
    if section_marker in existing:
        raise RuntimeError(f"{filename}: section already exists")
    page = doc.add_paragraph()
    page.add_run().add_break(WD_BREAK.PAGE)
    heading = doc.add_paragraph()
    run = heading.add_run(title)
    run.bold = True
    run.font.size = doc.styles["Normal"].font.size
    for paragraph in paragraphs:
        doc.add_paragraph(paragraph)
    doc.save(path)


append_section(
    "AI开发项目_项目说明文档.docx",
    "v1029｜微信首页视觉热修（2026-08-21）",
    [
        "在不合入并行开发中的 v1030 外卖结算等工作前提下，基于线上 v1029 单独更新微信第一个页面。聊天、联系人、群聊、朋友圈、角色回复、真实好友数据和外置语音配置均沿用原数据与逻辑。",
        "微信首页新增真实微信风格的夜间／白天双模式：标题固定并严格居中，无小耳朵；右侧加号缩为 29px；搜索框、Windows 微信已登录提示、聊天行灰阶和细分隔线按参考图调整。置顶聊天在夜间比普通行浅一档，白天用浅灰区分。",
        "顶部与底部导航均使用半透明渐变、backdrop-filter 模糊和饱和度形成磨砂透视。列表从固定顶部导航后方滚动，标题不随内容移动。底栏四个入口未选中为线框，选中为微信绿实心；不模拟消息红点。",
        "缓存热修标识为 v1029-wechat-home-1，Service Worker 使用独立 north-shell-v1029-wechat-home-1 缓存，避免线上继续返回旧首页。共享网页与私人 PhoneWeb.bundle 已同步。",
        "验证：干净 origin/main 隔离工作树 Node 自动复核 934/934 通过；390×844 浏览器实测标题中心、29px 加号、无耳朵／红点、置顶与普通行色差、顶部 36px 和底部 34px 磨砂均符合预期。Mac 编译、签名和真实 iPhone 验证未在本次完成。",
    ],
)

append_section(
    "AI开发项目_Bug记录模板.docx",
    "v1029｜微信首页视觉热修（2026-08-21）",
    [
        "问题：原微信首页标题和底栏仍使用通用组件，标签为“好友／朋友圈”，顶部会呈现普通实色导航，底栏选中态与真实微信不同；置顶行和普通行缺少默认色差。用户还明确要求移除标题旁小耳朵和全部模拟红点。",
        "根因：微信主页面复用了通用 nav／tabbar 和通用 SVG 图标，滚动层未延伸到固定导航后方，导致即使设置 backdrop-filter 也缺少真实透视层次；聊天行统一使用同一背景变量。",
        "处理：只重构 renderWeChat、wxChats 的页面框架和专属 SVG／CSS，保留所有真实数据映射与点击行为。顶部导航固定，滚动层以负边距和等量内边距从其后方经过；底栏按四种选中态切换线框／实心 SVG；置顶行使用独立背景；微信页面的原生状态栏主题跟随 wxTheme。",
        "缓存：不提升网页主版本，不带入并行 v1030；改用 v1029-wechat-home-1 Service Worker 热修标识和独立 shell cache，使线上 v1029 获取本次首页文件。",
        "验证结果：干净基线 934/934 通过；浏览器计算值确认标题中心 195.2px／视口中心 195px、加号 29px、耳朵 0、红点 0、夜间置顶 rgb(32,32,35)／普通 rgb(25,25,27)、顶部 blur(36px)／底部 blur(34px)。",
    ],
)

append_section(
    "AI开发项目_Bug修改规范.docx",
    "v1029｜微信主界面视觉层与数据层隔离规则（2026-08-21）",
    [
        "仿真社交软件首页时，视觉重构必须限定在外框、专属类名、图标和展示层；不得用静态假聊天替换正式联系人、群聊、消息排序、未读、朋友圈或点击动作。预览数据只能在 NORTH_PREVIEW 路径存在。",
        "固定磨砂导航不能只写透明色和 backdrop-filter。滚动内容必须真实延伸到导航后方，同时用固定 z-index 保证标题和按钮不移动，才能形成可验证的透视层次。",
        "底栏选中态应由同一数据状态驱动，未选中和选中 SVG 必须成对存在；用户明确不需要的红点不得根据其他数据自动模拟。人物实心图的头部与肩部应保留可见空隙。",
        "并行脏工作区需要单独上线时，必须从最新 origin/main 建立干净隔离工作树，只重放目标功能，更新独立缓存标识，完整测试并逐文件审查后再推送；禁止整文件暂存包含其他工作的混合文件。",
    ],
)

append_section(
    "AI开发项目_新聊天启动说明.docx",
    "v1029｜微信首页视觉热修交接（2026-08-21）",
    [
        "线上 v1029 额外包含微信首页视觉热修，缓存标识 v1029-wechat-home-1。该提交从干净 origin/main 隔离生成，不包含另一个工作区正在开发的 v1030 外卖结算、时区缓存或其他未提交内容。",
        "微信第一个页面已完成夜间／白天模式、固定居中标题、29px 加号、顶部和底部磨砂透视、搜索与 Windows 登录条、聊天灰阶、置顶行色差，以及四个底栏线框／绿色实心选中态。标题小耳朵和模拟消息红点均已移除。",
        "正式聊天、联系人、角色群、小手机真人好友／群聊、朋友圈、角色回复、后台通知和外置语音配置没有被替换或清空。私人 PhoneWeb.bundle 已由清单脚本同步同一实现。",
        "干净基线自动复核 934/934 通过，390×844 浏览器视觉和滚动坐标复核通过。尚未完成 Mac 编译、签名和真实 iPhone 验证；后续若继续改第二个页面，仍需在当前并行工作基础上谨慎合并本提交。",
    ],
)

print("Updated four maintenance DOCX files for v1029 WeChat home hotfix")

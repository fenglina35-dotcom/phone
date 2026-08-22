from pathlib import Path

from docx import Document


ROOT = Path(__file__).resolve().parents[1]
DOCS = ROOT / "docs" / "maintenance"

UPDATES = {
    "AI开发项目_项目说明文档": [
        "v1042｜跨场景记忆与真实外卖人工验证续跑（2026-08-22）",
        "微信、一次性线下约会与共同生活继续保留各自可见记录，但在“约会中同步到线上”开启时，共用按真实时间合并的隐藏剧情时间线。当前场景只认本场待回复输入；旁白保持事件身份，不得冒充角色台词。三个场景均可读取当前微信账号下该角色本人最近真实发布的朋友圈，只有相关时才自然承接。",
        "淘宝闪购真实外卖复用同一持久化浏览器与短时状态缓存，并按页面真实内容就绪提前继续。明确指定瑞幸、喜茶、霸王茶姬或奈雪时优先限定匹配品牌，不得换品牌兜底。",
        "平台弹出滑块、图片选择或其他真人验证时，不绕过平台风控。浏览器保持在前台，允许用户在两分钟内手动完成；验证消失后从原搜索页继续。自动搜索净耗时上限为35秒，角色任务运行态上限为3分钟，失败或超时后只汇报一次，必须等用户新的明确消息才能再次开始。",
        "普通微信正文、主模型失败后切换副模型、朋友圈评论和后台通知路线均未改动；真实外卖状态只约束外卖指令，禁止固定系统文字冒充角色回复。",
    ],
    "AI开发项目_Bug记录模板": [
        "v1042 Bug 记录｜淘宝闪购图片验证与角色重复搜索（2026-08-22）",
        "现象：淘宝闪购触发“请选择符合描述的所有图片”后，旧检测只识别主页面中的滑块／安全验证文字，无法识别跨框架图片验证；真实外卖失败回执生成角色回复时，通用外卖提示仍允许再次输出搜索标签，角色因此反复说“等一下给你找”并重新开一轮。",
        "根因：riskCheck 只读取主 frame 的 body 且只匹配四类短语；搜索提取循环中没有持续复核验证页。角色外卖状态仅保存在内存中的进行中锁，失败结果没有持久终态；回执模型仍收到“可以开始真实搜索”的通用规则。",
        "修复：汇总主页面与所有 frame 的可见文本，识别图片验证、验证码、滑块和访问频繁；验证出现时前置浏览器并等待用户手动处理，消失后原地续跑。搜索净耗时35秒、人工验证120秒、恢复后的角色运行态3分钟均有硬上限。每个角色持久记录本轮 running／completed／failed，回执生成及重开 App 后禁止自动重试；只有新用户消息晚于本轮结束时间才允许新一轮。",
        "边界：没有实现验证码识别、点击、规避或反检测；没有修改普通聊天模型请求、主副模型切换、朋友圈或后台通知。失败时仍由真实角色模型按人设自然汇报，系统不得写固定角色台词。",
    ],
    "AI开发项目_Bug修改规范": [
        "v1042｜受保护外部平台人工验证与有限重试规范（2026-08-22）",
        "浏览器自动化遇到平台滑块、图片选择、短信验证码或风控页时，只能检测、暂停、把窗口交给用户并在用户完成后继续；禁止自动识别、代点、绕过验证码或加入反检测规避。",
        "每个外部操作必须同时有净自动耗时上限、人工处理上限和跨重开的任务终止上限。超时或失败只能生成一个真实终态；不得由模型回执、定时器、页面重开或副模型再次触发同一指令。下一轮必须有终态之后的新用户明确消息。",
        "外部任务的进行中／终态提示只能约束对应功能标签，不得堵塞普通聊天；普通聊天的主模型失败后切换副模型仍按原路线执行。系统事实只供角色模型自然表达，不得用固定客服文字冒充角色。",
    ],
}


def append_docx(stem: str, paragraphs: list[str]) -> None:
    path = DOCS / f"{stem}.docx"
    document = Document(path)
    marker = paragraphs[0]
    if any(marker == p.text.strip() for p in document.paragraphs):
        return
    document.add_paragraph("")
    for paragraph in paragraphs:
        document.add_paragraph(paragraph)
    document.save(path)


def append_text(stem: str, paragraphs: list[str]) -> None:
    path = DOCS / f"{stem}.txt"
    current = path.read_text(encoding="utf-8")
    if paragraphs[0] in current:
        return
    path.write_text(current.rstrip() + "\n\n" + "\n".join(paragraphs) + "\n", encoding="utf-8")


for name, rows in UPDATES.items():
    append_docx(name, rows)
    append_text(name, rows)

print("Updated v1042 maintenance DOCX and searchable TXT companions")

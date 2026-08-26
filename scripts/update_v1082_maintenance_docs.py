"""Append the v1082 companion-delivery and exact-clock release to maintenance records."""

from pathlib import Path

from docx import Document
from docx.enum.text import WD_BREAK


ROOT = Path(__file__).resolve().parents[1]
DOCS = ROOT / "docs" / "maintenance"
MARKER = "v1082｜"


def append_docx(filename: str, title: str, paragraphs: list[str]) -> None:
    path = DOCS / filename
    doc = Document(path)
    existing = "\n".join(p.text for p in doc.paragraphs)
    if MARKER in existing:
        raise RuntimeError(f"{filename}: v1082 section already exists")
    page = doc.add_paragraph()
    page.add_run().add_break(WD_BREAK.PAGE)
    heading = doc.add_paragraph()
    run = heading.add_run(title)
    run.bold = True
    run.font.size = doc.styles["Normal"].font.size
    for paragraph in paragraphs:
        doc.add_paragraph(paragraph)
    doc.save(path)


def append_txt(filename: str, title: str, paragraphs: list[str]) -> None:
    path = DOCS / filename
    existing = path.read_text(encoding="utf-8")
    if MARKER in existing:
        raise RuntimeError(f"{filename}: v1082 section already exists")
    block = "\n\n" + title + "\n" + "\n".join(paragraphs) + "\n"
    path.write_text(existing.rstrip() + block, encoding="utf-8", newline="")


records = [
    (
        "AI开发项目_项目说明文档",
        "v1082｜共同生活伴生畅通与真实时间复核（2026-08-27）",
        [
            "版本：共享网页 v1082；私人 iOS 1.0.207（207）；原生桥 25。v1081 的自然点单、KFC 首页套餐、真实外卖回执和此前功能全部保留。",
            "共同生活开启期间，伴生后台的普通文字消息不再被服务器暂停，也不再在客户端收件箱积压。服务器已生成的消息会立即写入对应线上聊天并持久化确认；本机伴生自动消息采用同一条共同生活不阻塞的通道。关闭共同生活后不会再集中涌出旧消息。",
            "共同生活状态仍会作为实时事实进入角色上下文：角色知道双方正在同居或面对面，不能把后台文字说成异地、久未见、等待落地或催用户回复。来电、一次性线下约会、角色忙碌和放映室等原有占用边界没有因文字畅通而放宽。",
            "每次微信角色回复开始生成前，程序都从真实后台时钟 Date.now() 重新取得当前时刻，再按角色时区换算并作为本轮最后一条内部时间事实；不读取聊天上下文里的旧时间。生成结束、真正发送前会再次读取最新后台时间复核。明确当前时间误差超过两分钟时只允许一次真实模型重写；失败保持无伪造回复。与时间无关时不要求角色报时。",
            "最终时间校验覆盖无前缀的当前整点说法，例如真实 03:50 时“4点了，快睡”会被拦截；03:50 说 03:52 在允许误差内；“已经约好四点见”“等到四点了再叫我”等日程表达不会被误判。",
        ],
    ),
    (
        "AI开发项目_Bug修改规范",
        "v1082｜伴生消息送达与后台时间规范（2026-08-27）",
        [
            "共同生活状态不能作为伴生后台文字消息的服务器暂停条件或客户端收件阻塞条件。已生成消息必须先成功写入并持久化，再确认服务器收件；不得留在队列中等待用户关闭共同生活。",
            "文字畅通与场景事实必须分离：共同生活期间允许伴生文字送到线上聊天，但角色仍必须知道双方正在共同生活，不能虚构异地、分居、久未见、等待落地或催回复。此规则不授权后台自行来电。",
            "角色当前时间必须来自本轮实际读取的后台时钟，不得沿用系统提示、聊天记录、旧消息时间戳或模型猜测。初次生成前和最终发送前都要读取最新值，并按角色时区换算。",
            "当前几点或经过多久的误差在一至两分钟内允许自然表达；超过两分钟、提前跳整点或跨小时必须在发送前拦截，并只做一次真实模型纠正。纠正失败不得用固定台词冒充角色。日程、约定和未来时间不能被误判为当前时间。",
        ],
    ),
    (
        "AI开发项目_Bug记录模板",
        "v1082｜共同生活积压伴生消息与角色跳时（2026-08-27）",
        [
            "现象：开启共同生活后，伴生后台消息不能进入线上聊天；关闭后积压消息一次性涌出。真实 03:50 时，角色可能说成 04:00，只有用户要求看时间后才纠正。",
            "根因：服务器自动任务配置、客户端伴生自动触发和服务器消息补拉共同复用了包含共同生活静默的普通主动消息门禁，导致生成或送达被延后。时间提示虽已读取后台时钟，但位置不够靠近当前生成，且最终校验按标点拆句后漏掉无“现在”前缀的“4点了”。",
            "修复：为伴生后台建立不受共同生活影响、但仍尊重一次性约会/忙碌/通话等边界的送达门禁；服务器配置、本机伴生自动消息和补拉统一使用该门禁。每次回复追加最新后台时间钉，并把唯一一次时间纠正移到所有正文重写之后、消息解析发送之前。",
            "验证：根目录全量 1090/1090、外卖服务 154/154 通过；网页和私人 iOS app.js SHA-256 一致。专项覆盖共同生活服务器生成不暂停、客户端不积压、03:50 拦截“4点了”、两分钟容差和日程不误伤。",
        ],
    ),
    (
        "AI开发项目_新聊天启动说明",
        "v1082｜新聊天接手说明（2026-08-27）",
        [
            "当前候选：网页 v1082；私人 iOS 1.0.207（207）；原生桥 25。先核对远端提交、ZIP SHA-256、根/私人 app.js 与 delivery.js 哈希，并保护大量未跟踪用户资料。",
            "共同生活期间伴生后台普通文字应即时进入线上聊天，不能积压到关闭后集中出现。角色仍知道正在共同生活，不得说成异地、久未见或催回复；共同生活期间后台来电仍受限制。",
            "角色每次微信回复前和最终发送前都使用真实后台 Date.now() 的最新值并按角色时区换算；不是聊天上下文时间。允许一至两分钟自然误差，超过两分钟或跳整点必须由真实模型纠正一次，与时间无关时不主动报时。",
            "KFC 九珍果汁饮料/爆汁三柠茶真实规格选择仍未在营业门店验收，不能写成已完成。Windows 测试和 ZIP 结构校验不能写成 Mac 编译、签名或真实 iPhone 已通过。",
        ],
    ),
]

for stem, title, paragraphs in records:
    append_docx(stem + ".docx", title, paragraphs)
    append_txt(stem + ".txt", title, paragraphs)

print("Updated four maintenance DOCX/TXT pairs for v1082")

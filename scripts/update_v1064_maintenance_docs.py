from pathlib import Path

from docx import Document


ROOT = Path(__file__).resolve().parents[1]
DOCS = ROOT / "docs/maintenance"
MARKER = "v1064／1.0.187"


def append_section(filename, title, paragraphs):
    path = DOCS / filename
    document = Document(path)
    if any(MARKER in paragraph.text for paragraph in document.paragraphs):
        print(f"SKIP={path.name}")
        return
    document.add_heading(title, level=1)
    for paragraph in paragraphs:
        document.add_paragraph(paragraph)
    document.save(path)
    print(f"UPDATED={path.name}")


append_section(
    "AI开发项目_项目说明文档.docx",
    "v1064／1.0.187 外卖澄清续单与网关恢复（2026-08-25）",
    [
        "网页核心升级为 v1064，私人 iOS 升级为 1.0.187（187），原生桥继续为 25。真实复测确认：主商品成功加购后平台购物车底栏会短暂重绘，旧流程在重绘完成前读取到陈旧结算控件，因而提前尝试结算并跳过同 taskId 的后续商品。",
        "修复后先被动等待并重新读取权威的起送状态，再续接原 taskId、原门店和原主商品执行 addRequestedStandaloneItems。起送短缺金额与起送门槛分开解析；在页面显示购物车 10.9 元、还差 9.1 元时，门槛计算为 20 元。",
        "真实浏览器验证已完成但未付款：杨姥佬家de撒汤门店中，（招牌）营养鸡丝撒汤、爆款鲜肉煎饺（8个）、金陵/灌汤包（6个）各一份，总额 24.5 元，进入订单确认页。没有重复主商品，没有新建第二个任务。",
        "Supabase phone-delivery 对无具体业务正文的中间层 HTTP 502 仅允许一次原请求重试，create_order 在允许列表中，pay_order 不允许自动重试；具体浏览器或平台错误保持原样返回。普通角色回复、朋友圈、后台通知、远控字幕及外置语音配置未改路由。Windows 自动测试已通过；Mac 编译、签名和真实 iPhone 覆盖安装仍待执行。",
    ],
)

append_section(
    "AI开发项目_Bug记录模板.docx",
    "v1064／1.0.187 外卖澄清续单与网关恢复 Bug 记录（2026-08-25）",
    [
        "现象：撒汤已加入购物车，角色询问煎饺或灌汤包后用户回答“都要”，自动化没有补商品，或者提示购物车有商品但未进入确认页；另有偶发“真实外卖上游 HTTP 502”。",
        "根因一：平台购物车关闭后的底栏异步重绘，旧逻辑立即读取并误点陈旧 checkout 控件。根因二：短缺文案“差 9.1 元起送”被旧正则误当成明确起送门槛。根因三：边缘函数对中间层无正文 502 没有有界恢复。",
        "修复：购物车恢复后最多七次、每次 250 毫秒被动重读，不进行额外搜索或重复加购；发现权威短缺文案后续接原任务补单。门槛由可见购物车金额加短缺金额推导。通用 502 只重试一次，支付动作永不自动重试。",
        "验证：真实浏览器同 taskId 第二次修订成功补齐撒汤、煎饺和灌汤包，各一份并进入订单确认页；未提交、未支付。自动测试同时覆盖现实页面金额解析、被动重读、通用 502 单次重试及 pay_order 排除。",
    ],
)

append_section(
    "AI开发项目_Bug修改规范.docx",
    "v1064／1.0.187 外卖续单与网关恢复规范（2026-08-25）",
    [
        "续单规范：用户对起送价澄清的回答必须绑定原 taskId，保留原门店、原主商品和完成清单；不得重新执行首单意图，不得重复添加已完成商品。",
        "页面规范：平台加购、关闭购物车或切换规格后可能异步重绘。自动化必须先做有上限的被动重读，只有权威页面状态稳定后才能继续加购或结算；被动重读不得转化成新的搜索。",
        "金额规范：短缺金额不是起送门槛。存在购物车实付和短缺时，门槛为二者之和；没有可靠金额时应暂停说明，不能伪造阈值。",
        "网关规范：仅对通用、无具体业务错误的 HTTP 502 做一次同体同 taskId 重试；支付、取消和其他不可安全重复动作不得自动重试。发布前必须验证普通回复、朋友圈、后台通知、远控字幕和外置语音配置未受影响。",
    ],
)

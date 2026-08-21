from pathlib import Path

from docx import Document
from docx.enum.text import WD_BREAK


ROOT = Path(__file__).resolve().parents[1]
DOCS = ROOT / "docs" / "maintenance"
ZIP_NAME = "SmallPhone_v1031_WeChatHome_iOS152_MacReady.zip"
ZIP_SHA256 = "99030991a41c420b4f531c49c2492fdfd29365a05b2888051d2090c4ed314195"


def append_section(filename, title, paragraphs):
    path = DOCS / filename
    doc = Document(path)
    existing = "\n".join(p.text for p in doc.paragraphs)
    marker = title.split("（", 1)[0]
    if marker in existing:
        raise RuntimeError(f"{filename}: section already exists")
    page = doc.add_paragraph()
    page.add_run().add_break(WD_BREAK.PAGE)
    heading = doc.add_paragraph()
    run = heading.add_run(title)
    run.bold = True
    run.font.size = doc.styles["Normal"].font.size
    for paragraph in paragraphs:
        doc.add_paragraph(paragraph)
    doc.save(path)


append_section(
    "AI开发项目_项目说明文档.docx",
    "v1031｜私人 iOS 1.0.152 安装包（2026-08-21）",
    [
        "为使私人 App 包含已上线的 v1031 微信首页，私人 iOS 从 1.0.150（150）升为 1.0.152（152），原生桥保持 25。1.0.151 已由并行 v1030 支付回执工作占用，因此不重复使用。",
        "私人 PhoneWeb.bundle 内嵌网页为 v1031，包含 v1029 真实外卖与角色钱包基线及微信首页视觉升级；不包含未合入的 v1030 支付回执闭环。",
        f"交付物仅一个 ZIP：{ZIP_NAME}。包内 176 个文件，只有一份 PhoneWeb.bundle/index.html，只保留 Mac 编译前说明和第一百五十二次安装说明，无嵌套 ZIP 或 Python 缓存。SHA-256：{ZIP_SHA256}。",
        "Windows Node 自动复核 934/934 通过。Mac 编译、签名和真实 iPhone 覆盖安装仍未完成。",
    ],
)

append_section(
    "AI开发项目_Bug记录模板.docx",
    "v1031｜其他工作区打包可能漏掉微信新版（2026-08-21）",
    [
        "问题：用户担心由另一工作区直接打私人包时，是否会自动带入 v1031 微信首页。",
        "根因：Git 工作树隔离只共享对象库，不会自动把一个分支的提交应用到另一工作树。另一工作区若没有同步 main 或合入 v1031，打包必然漏掉本次微信修改。",
        "处理：直接从干净 v1031 隔离工作树升级私人 iOS 1.0.152（152）并制作唯一 MacReady ZIP；他人后续打包时必须先同步已包含 v1031 的 main。",
        f"验证：针对性版本契约 69/69 通过，最终全量 934/934 通过；ZIP 文件数 176，SHA-256 {ZIP_SHA256}。",
    ],
)

append_section(
    "AI开发项目_Bug修改规范.docx",
    "v1031｜跨工作树私人包合入规则（2026-08-21）",
    [
        "私人包只会包含打包工作树当前实际文件，不会因另一分支已提交或已推送而自动合入。打包人必须先确认 HEAD 已包含目标提交，再重建 PhoneWeb.bundle。",
        "并行版本已占用私人 iOS build 时，新交付使用下一个未占用 build，不得为两套不同源码发布相同 build。本次因 151 已占用而使用 152。",
        "交付目录只保留一个最终 ZIP；必须在压缩后再检查网页版本、iOS marketing/build、PhoneWeb.bundle/index.html 数量、说明文件数量、嵌套 ZIP、临时缓存、文件数与 SHA-256。",
    ],
)

append_section(
    "AI开发项目_新聊天启动说明.docx",
    "v1031｜私人 iOS 1.0.152 交付交接（2026-08-21）",
    [
        "当前交付版本：网页 v1031；私人 iOS 1.0.152（152）；原生桥 25。远端 main 已包含 v1031 微信首页；私人安装包从同一干净基线生成。",
        f"唯一交付物：delivery-v1031/{ZIP_NAME}；176 个文件；SHA-256 {ZIP_SHA256}。",
        "本包不含并行 v1030 支付回执闭环。其他工作区若要制作后续私人包，必须先同步包含 v1031 的 main，并使用高于 v1031／1.0.152 的新版本，不得降版。",
        "Windows 自动复核 934/934 通过。Mac 编译、签名和真实 iPhone 覆盖安装未完成。",
    ],
)

print("Updated four maintenance DOCX files for v1031 private package")

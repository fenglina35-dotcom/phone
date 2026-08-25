from pathlib import Path

from docx import Document


ROOT = Path(__file__).resolve().parents[1]
DOCS = ROOT / "docs/maintenance"
MARKER = "v1065／1.0.188"


def append_section(filename, title, paragraphs):
    path = DOCS / filename
    document = Document(path)
    if any(MARKER in paragraph.text for paragraph in document.paragraphs):
        print(f"SKIP={path.name}")
        return
    document.add_heading(title, level=1)
    for paragraph in paragraphs:
        document.add_paragraph(paragraph)
    document.save(path)
    print(f"UPDATED={path.name}")


append_section(
    "AI开发项目_项目说明文档.docx",
    "v1065／1.0.188 外卖澄清修订连续性修复（2026-08-25）",
    [
        "网页核心升级为 v1065，私人 iOS 升级为 1.0.188（188），原生桥继续为 25。本次修复真实聊天中的同任务续单：初始搜索使用修订 1，系统发出起送价澄清时增加到修订 2，用户回答“都要”等澄清内容后必须沿用修订 2继续，不能再次增加到 3。",
        "此前的 502 不是平台搜索失败，也不是 Cloudflare 隧道离线，而是前端把同一个 taskId 的修订从 1 跳到 3，浏览器适配器按连续性守卫拒绝该请求；本地服务与 Edge 层又把任务状态冲突错误归类成通用上游 502。",
        "修复后任务、修订、澄清、约束和状态冲突统一保留为 HTTP 409，只有真正的上游故障才显示 502。可执行测试直接运行真实 delivery.js，覆盖撒汤首次因起送价不足、角色给出煎饺与灌汤包、用户回答“都要”的完整聊天路径，断言修订严格为 1、2，三项商品各一份，并停在待支付。",
        "根网页与私人 PhoneWeb.bundle 的 delivery.js 同步；普通角色回复、朋友圈回复、后台通知、远控字幕和外置语音配置不改路由。Windows 自动测试覆盖发布边界；Mac 编译、签名和真实 iPhone 安装仍需在 Mac 上完成。",
    ],
)

append_section(
    "AI开发项目_Bug记录模板.docx",
    "v1065／1.0.188 外卖澄清修订连续性 Bug 记录（2026-08-25）",
    [
        "现象：撒汤已找到并因起送价不足询问补煎饺还是灌汤包，用户回答“都要”且角色明确同意后，自动化没有继续，并显示“真实外卖上游 HTTP 502”。",
        "根因：requestRoleClarification 已把同一任务修订从 1 增为 2，澄清回答进入 roleRequest 的换品分支时又错误增为 3；适配器只接受当前修订或连续的 +1，因此拒绝 1→3。拒绝信息又被本地服务和 Edge 函数错误映射成 502。",
        "修复：换品/澄清回答只清除等待状态并更新 orderIntent，不再增加修订；回答继续绑定原 taskId、门店、主商品和完成清单。任务状态冲突返回 409，禁止伪装成上游网络错误。",
        "回归：真实 delivery.js 虚拟机测试断言请求修订仅为 [1, 2]，初次失败草稿后只续单一次，撒汤、煎饺、灌汤包各一份，无重复商品、不新建任务、不付款。",
    ],
)

append_section(
    "AI开发项目_Bug修改规范.docx",
    "v1065／1.0.188 外卖任务修订与错误分类规范（2026-08-25）",
    [
        "修订规范：创建任务时修订为 1；每次系统新发出一个澄清问题只增加一次；用户回答该问题时沿用当前修订，不得再次增加。适配器应接受同修订的幂等重试或严格 +1 的新阶段，拒绝跳级。",
        "续单规范：澄清回答必须复用原 taskId、角色、账号、会话、门店、商品进度和已完成清单；不得重新执行 roleRequestIntent，不得重复主商品，不得从旧消息或页面恢复中创建新任务。",
        "错误规范：任务、修订、澄清、约束和状态冲突属于 HTTP 409；只有真实上游无响应或网关故障才是 502。界面必须保留可执行原因，不能把程序状态错误说成平台搜索不到。",
        "发布规范：至少测试同 taskId 幂等、完成/过期/取消不可恢复、普通聊天不触发、当前回合结构化动作可触发、澄清不新建任务，以及普通回复、朋友圈、后台通知、远控字幕和外置语音配置隔离。",
    ],
)

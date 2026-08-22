from pathlib import Path
from docx import Document


ROOT = Path(__file__).resolve().parents[1]
DOCS = ROOT / "docs" / "maintenance"

UPDATES = {
    "AI开发项目_项目说明文档": [
        "v1043｜微信小号搜索与真实外卖低验证直达（2026-08-22）",
        "共享网页升级为v1043，私人iOS升级为1.0.162（162），原生桥保持25。小号可在新的朋友页面按角色名字、微信号或手机号搜索并添加，手机号兼容空格、横杠和中英文括号；添加关系只写入当前小号。白色主题空结果加号使用白色微信表面和绿色符号；角色／对方自定义气泡不再强制增加高对比描边，用户侧气泡可读性描边保持不变。",
        "真实外卖浏览器在第一次成功定位并读取真实商品规格后保存受控本机店铺／商品入口。同款即使糖度、冰度等偏好变化也优先直达，并重新读取当前价格、库存和规格；失效或超过30天才回到全平台搜索。平台强制登录、验证码、风控和支付验证仍须本人处理，不绕过平台。",
        "普通微信回复、主模型失败后切换副模型、朋友圈、后台通知、外置语音配置和既有引用／视频／游戏修复均保持原路线。",
    ],
    "AI开发项目_Bug记录模板": [
        "v1043 Bug 记录｜微信小号角色搜索缺失与同款外卖重复验证（2026-08-22）",
        "现象：小号新的朋友只能筛好友申请，不能用角色微信号、手机号或名字找到并添加角色；白色主题空状态加号底色和角色气泡描边不符合界面设定。同一真实饮品每次都重新全平台搜索，容易反复进入图片验证。",
        "修复：增加统一联系人搜索标准化和当前小号独立添加；补白色主题空状态样式并仅取消角色／对方气泡强制描边。真实外卖保存首次成功读取规格的店铺商品入口，直达后重读当次数据；临时网络失败不删除路线，入口失效才清除，直达遇验证不启动第二轮搜索。",
        "验证：微信专项8/8、真实外卖服务18/18通过；完整网页自动化在发布前重新执行。网页和私人内置网页由同一清单重建并校验关键函数。Mac编译、签名和真实iPhone覆盖安装仍待完成。",
    ],
    "AI开发项目_Bug修改规范": [
        "v1043｜多工作区网页／私人包一致性与外部路线缓存规范（2026-08-22）",
        "共享网页和私人PhoneWeb.bundle同时存在改动时，发布前必须从共享清单重建私人包并执行私人转换，再逐项校验关键函数和版本号；不得用简单文件覆盖破坏私人专用转换。",
        "受控本机只允许缓存本人已经成功走通的真实商家／商品入口，不缓存价格、库存、优惠和规格。每次直达必须重读当前页面；验证出现时停止一次，禁止为了规避验证改入口或自动代点。",
    ],
}


def append_docx(stem, rows):
    path = DOCS / f"{stem}.docx"
    document = Document(path)
    if any(rows[0] == paragraph.text.strip() for paragraph in document.paragraphs):
        return
    document.add_paragraph("")
    for row in rows:
        document.add_paragraph(row)
    document.save(path)


def append_text(stem, rows):
    path = DOCS / f"{stem}.txt"
    current = path.read_text(encoding="utf-8")
    if rows[0] in current:
        return
    path.write_text(current.rstrip() + "\n\n" + "\n".join(rows) + "\n", encoding="utf-8")


for name, rows in UPDATES.items():
    append_docx(name, rows)
    append_text(name, rows)

print("Updated v1043 maintenance DOCX and TXT companions")

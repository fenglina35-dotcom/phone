"""Append the v1079 photo-card and points-coupon release to all maintenance records."""

from pathlib import Path

from docx import Document
from docx.enum.text import WD_BREAK


ROOT = Path(__file__).resolve().parents[1]
DOCS = ROOT / "docs" / "maintenance"
MARKER = "v1079｜"


def append_docx(filename: str, title: str, paragraphs: list[str]) -> None:
    path = DOCS / filename
    doc = Document(path)
    existing = "\n".join(p.text for p in doc.paragraphs)
    if MARKER in existing:
        raise RuntimeError(f"{filename}: v1079 section already exists")
    page = doc.add_paragraph()
    page.add_run().add_break(WD_BREAK.PAGE)
    heading = doc.add_paragraph()
    run = heading.add_run(title)
    run.bold = True
    run.font.size = doc.styles["Normal"].font.size
    for paragraph in paragraphs:
        doc.add_paragraph(paragraph)
    doc.save(path)


def append_txt(filename: str, title: str, paragraphs: list[str]) -> None:
    path = DOCS / filename
    existing = path.read_text(encoding="utf-8")
    if MARKER in existing:
        raise RuntimeError(f"{filename}: v1079 section already exists")
    block = "\n\n" + title + "\n" + "\n".join(paragraphs) + "\n"
    path.write_text(existing.rstrip() + block, encoding="utf-8", newline="")


records = [
    (
        "AI开发项目_项目说明文档",
        "v1079｜图文照片卡与吃货豆红包兑换闭环（2026-08-27）",
        [
            "版本：共享网页 v1079；私人 iOS 1.0.204（204）；原生桥 25。网页、私人 PhoneWeb.bundle、缓存号和 Xcode 构建号必须成套一致。",
            "未开启或未完整配置图片生成模型时，角色输出 [图片|画面描述] 会成为白色图文照片卡；配置有效时继续走原真实生图。卡片使用 42px 浅灰标题区和浅灰正文，正文高度自动增长、无最大高度和行数截断，长描述通过聊天区正常滚动完整查看。",
            "微信功能面板新增“图文描述”。用户填写的描述以 image/textCard 消息入库，角色把它当作图片事实理解并自然回应，只能使用描述中明确出现的视觉细节。",
            "真实外卖吃货豆红包按平台真实顺序执行：选中红包并点底部确定，出现“是否兑换”后点“立即兑换”，必要时再次确认并回到结算页核对红包和总价。兑换失败、页面异常或结果不确定时 fail-closed，禁止按原价继续提交。",
            "真实验收使用 200 吃货豆兑换 ¥5 红包，结算价由 ¥24.79 变为 ¥19.79，状态为 applied / coupon_selected_and_total_verified；最后仍在有“立即支付”的结算页，没有提交订单、没有付款。",
        ],
    ),
    (
        "AI开发项目_Bug修改规范",
        "v1079｜描述卡完整显示与积分红包确认规范（2026-08-27）",
        [
            "图文照片卡不得用固定高度、max-height、line-clamp、ellipsis 或隐藏溢出截断正文。标题可以固定高度，但正文必须参与卡片自动布局；超出当前屏幕时由聊天容器滚动，不能让后半段文字不可见。",
            "图文描述是图片事实，不是普通文本占位。模型历史必须明确 image/textCard 语义，并限制角色只依据描述回应；有真实生图配置时不得降级成图文卡。",
            "积分兑换券只有在真实卡片明确显示未兑换、且真实兑换弹窗出现时才能点击兑换。必须优先点精确 DOM 控件，不能误点吸底栏或装饰子元素；没有弹窗不得猜测兑换。",
            "红包兑换、选中、返回结算和总价核验是同一事务。任一步失败、状态不确定或平台报错都必须停止，绝不能为了完成点单而按原价继续进入提交或支付。",
        ],
    ),
    (
        "AI开发项目_Bug记录模板",
        "v1079｜无生图配置照片退化与吃货豆红包未闭环（2026-08-27）",
        [
            "现象：无图片生成配置时角色的照片只显示普通 [图片] 文字，用户也没有主动发送图文照片的入口；初版大号“照片”标题占用过多空间，且需要明确排除长描述被裁剪。真实外卖选择需吃货豆兑换的红包后，页面可能先出现临时选中，再弹兑换确认，旧流程可能错过弹窗或误点被吸底栏遮挡的控件。",
            "根因：图片解析器把无生图路线降级为 text；消息历史没有 textCard 图片语义；卡片样式未明确自动高度边界。外卖流程把“选中券”当作终态，没有把延迟出现的吃货豆兑换弹窗、再次确认和结算总价复核建成完整状态机。",
            "修复：新增 image/textCard 消息、角色与用户双向语义、图文描述入口和黑白主题白卡样式；标题缩至 14px/42px，正文浅灰且 height:auto、max-height:none、overflow:visible、line-clamp:unset。外卖新增积分券识别、延迟弹窗观察、精确 DOM 点击、兑换后再次确认、页面异常与不确定状态 fail-closed。",
            "验证：图文照片卡专项覆盖有配置真实生图、无配置卡片、用户图片语义、三份 HTML 黑白样式、长文本无截断和根/私人 app.js 一致。外卖服务 154 项与网页全量测试在发布前重新执行。真实系统已完成 200 吃货豆换 ¥5，¥24.79 → ¥19.79，未下单未付款。Mac 编译、签名和真实 iPhone 仍待用户完成。",
        ],
    ),
    (
        "AI开发项目_新聊天启动说明",
        "v1079｜新聊天接手说明（2026-08-27）",
        [
            "当前候选：网页 v1079；私人 iOS 1.0.204（204）；原生桥 25。先核对远端提交、ZIP SHA-256、根/私人 app.js 哈希和大量未跟踪资料保护状态。",
            "图文照片卡只有在图片生成未开启或配置不可用时作为角色回退；配置有效时必须继续真实生图。用户“图文描述”始终作为图片事实入聊。长描述不得截断。",
            "吃货豆红包真实验收已成功抵扣 ¥5，但没有下单或付款。后续若平台流程变化，必须保持兑换失败或状态不确定时停止，最终付款仍只由用户本人完成。",
            "Windows 自动测试和 ZIP 结构校验不能写成 Mac 编译、签名或真实 iPhone 已通过；普通角色回复、朋友圈、后台通知、远控字幕和外置语音配置继续作为发布保护线。",
        ],
    ),
]

for stem, title, paragraphs in records:
    append_docx(stem + ".docx", title, paragraphs)
    append_txt(stem + ".txt", title, paragraphs)

print("Updated four maintenance DOCX/TXT pairs for v1079")

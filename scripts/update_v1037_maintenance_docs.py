from pathlib import Path

from docx import Document
from docx.enum.text import WD_BREAK


ROOT = Path(__file__).resolve().parents[1]
DOCS = ROOT / "docs" / "maintenance"
MARKER = "v1037｜"


def append_section(filename, title, paragraphs):
    path = DOCS / filename
    doc = Document(path)
    existing = "\n".join(p.text for p in doc.paragraphs)
    if MARKER in existing:
        for paragraph in doc.paragraphs:
            if "943/943" in paragraph.text:
                paragraph.text = paragraph.text.replace("943/943", "944/944")
            if paragraph.text == "v1037｜微信个人页、好友交互与输入稳定性完善（2026-08-22）":
                paragraph.text = "v1037｜微信个人页、好友交互、真实外卖订单与输入稳定性完善（2026-08-22）"
            if paragraph.text.startswith("角色可见回复继续来自真实模型") and "真实外卖订单聊天卡片" not in paragraph.text:
                paragraph.text += " 真实外卖订单聊天卡片只展示连接器回执的商家、规格、金额、优惠、本人付款入口和配送状态，失败时不伪造订单。"
            if paragraph.text.startswith("下一步现实动作") and "真实外卖订单卡片" not in paragraph.text:
                paragraph.text += " 同时验证真实外卖订单卡片、本人付款入口和配送状态刷新。"
        doc.save(path)
        return
    page = doc.add_paragraph()
    page.add_run().add_break(WD_BREAK.PAGE)
    heading = doc.add_paragraph()
    run = heading.add_run(title)
    run.bold = True
    run.font.size = doc.styles["Normal"].font.size
    for paragraph in paragraphs:
        doc.add_paragraph(paragraph)
    doc.save(path)


append_section(
    "AI开发项目_项目说明文档.docx",
    "v1037｜微信个人页、好友交互与输入稳定性完善（2026-08-22）",
    [
        "版本：共享网页 v1037；私人 iOS 1.0.157（157）；原生桥 25。网页与私人 App 的 app.js、wechat-me.js、wechat-me.css、二维码库、缓存身份、资源清单和版本显示已经同步。",
        "本版完成微信“我”页及独立子页面：个人资料与人设、真实可扫描的好友二维码与扫一扫、服务、模拟钱包/零钱/银行卡/亲属卡/账单、实用型小手机客服、收藏、朋友圈相册、表情开发占位、微信独立设置和多角色账号切换。未加入退出账号、微信豆、发票、经营账户、支付分、小程序和小店卡包。",
        "收藏支持文字、语音、图片和翻译内容，支持重复播放与取消收藏；朋友圈相册删除图片时同步删除对应朋友圈。微信加号菜单恢复扫一扫；听一听与搜一搜返回时仍停留在微信发现页。火车票入口复用现有云程页面。",
        "聊天顶部心情气泡与“让ta回”操作保持，原先夹在聊天区与输入区之间的整条分隔背景已移除。附近好友刷新可生成不同性格与风险倾向的人，所有匿名新好友统一使用灰色头像。小手机群聊与新建群聊改为跟随微信深浅主题，浅色采用有层次的灰白底，避免与微信背景融为一体。",
        "共同生活输入框已撤销会替换 textarea 的光标修补，恢复原生 textarea 键盘链路，只保留非侵入式 caret-color。没有新增 focus、replaceWith、appearance 或固定高度等会遮挡键盘、阻断输入的处理。",
        "全量 Node 自动复核 944/944 通过，并已把同一套资源重新打入 PhoneWeb.bundle。Windows 不能替代 Mac 编译、签名、覆盖安装和真实 iPhone 触摸、键盘、相机扫码、语音播放、滚动及性能验证。",
        "角色可见回复继续来自真实模型与既有人设，不用固定角色台词；朋友圈模型回复、普通聊天回复、后台通知与外置语音配置边界保持不变。钱包、银行卡和亲属卡均为小手机内部模拟资产，不连接真实银行或真实微信支付。",
    ],
)

append_section(
    "AI开发项目_Bug记录模板.docx",
    "v1037｜共同生活键盘、微信分隔条与好友头像修复（2026-08-22）",
    [
        "现象：共同生活输入框为了改变光标曾把原生 textarea 替换成 input，并加入焦点和尺寸修补，导致 iOS 键盘遮挡、输入框无法输入。微信旧好友与新好友聊天页还出现与背景同色但占据高度的整条分隔区域；附近好友与新朋友头像样式不统一。",
        "根因：输入控件类型和焦点生命周期被视觉修补改变；聊天输入区的手动回复容器仍以整行背景参与布局；匿名好友头像没有走统一的占位渲染路径。",
        "处理：恢复共同生活原生 textarea 与旧自动高度行为，仅设置 caret-color；移除替换控件、强制 focus、appearance、touch-action 和固定行高等逻辑。聊天心情气泡和“让ta回”按钮不删除，只把其外层整行容器改为透明、无边框且不占满宽度。匿名新好友统一输出灰色头像组件。",
        "相邻完善：微信“我”页整套独立页面、收藏与取消收藏、朋友圈相册联动删除、扫一扫、微信设置、多账号切换、附近好友混合性格刷新、群聊主题跟随以及发现页返回路径均已加入专项测试。",
        "验证：微信与键盘专项测试通过；共享网页与私人 PhoneWeb.bundle 完成一致性同步；全量 944/944 通过；git diff --check 无空白错误。尚未完成 Mac 编译签名和真实 iPhone 键盘、扫码、语音、长列表滚动及覆盖安装验证。",
    ],
)

append_section(
    "AI开发项目_Bug修改规范.docx",
    "v1037｜移动输入控件与微信页面回归规范（2026-08-22）",
    [
        "移动端输入问题不得通过替换 textarea/input 元素、强制重新聚焦或固定输入区高度来修复视觉光标。优先保持原控件类型、事件绑定、selection、自动高度和系统键盘生命周期；纯视觉需求只能采用不改变交互语义的 CSS。",
        "聊天页删除分隔背景时必须区分内容控件与承载层：保留心情气泡、回复按钮及其点击能力，只让不需要的整行容器透明、无边框、非满宽。旧好友、新好友、角色聊天和小手机好友聊天必须分别回归。",
        "匿名好友占位头像必须走统一组件，不能因入口不同而显示 emoji、首字母或不同颜色。附近好友的性格多样化只改变虚拟资料和对话提示，不得绕过骚扰拦截、拉黑、举报或隐私边界。",
        "微信深浅主题应作为群聊与新建群聊的全局基底；单角色气泡设置优先于全局气泡设置。浅色主题必须保留容器层次和边界，避免纯白页面与卡片完全融合。",
        "新增微信独立页面时必须维持调用栈返回：从发现进入的听一听、搜一搜或云程应返回发现，从主页进入才返回主页。不得把微信设置与 API、模型或外置语音密钥混在同一页面。",
        "任何涉及角色回复的改动都必须保留真实模型生成、普通回复、朋友圈回复、后台通知和外置语音配置；失败时不得伪造角色台词或用系统固定文本冒充角色。",
    ],
)

append_section(
    "AI开发项目_新聊天启动说明.docx",
    "v1037｜微信个人页与好友交互完善交接（2026-08-22）",
    [
        "当前候选基线：网页 v1037；私人 iOS 1.0.157（157）；原生桥 25；分支 main。共同生活输入框已恢复原生 textarea，微信“我”页整套页面、二维码/扫一扫、服务与模拟钱包、收藏、朋友圈相册、微信设置和账号切换均已完成。",
        "附近好友支持混合性格刷新和统一灰色匿名头像；聊天心情气泡与“让ta回”保留，整条分隔背景移除；小手机群聊和新建群聊随微信黑白主题变化。听一听、搜一搜和云程返回路径保持在微信内部。",
        "自动复核 944/944 通过，私人 PhoneWeb.bundle 已重新同步。Windows 未执行 Mac 编译、签名、覆盖安装或真实 iPhone 验证；相机扫码、系统键盘、语音重复播放、长列表滚动和前后台状态仍需在真机完成。",
        "交付规则：每次只生成并发送一个最新 ZIP。v1037 包应只有一个顶层目录、完整 Xcode 工程、一个 PhoneWeb.bundle/index.html、无嵌套 ZIP、无缓存预览和旧安装说明；根部只保留请在Mac编译前先读.md与第一百五十七次安装说明。",
        "下一步现实动作：在 Mac 全新目录解压，打开 PhoneCompanionTest.xcodeproj，选择自己的签名团队后覆盖安装且不要卸载旧 App；核对网页 v1037、私人 iOS 1.0.157（157）和原生桥 25，再按安装说明验证共同生活输入、扫码、收藏、相册删除、附近好友、群聊主题、返回路径与滚动稳定性。",
    ],
)

print("Updated four maintenance DOCX files for v1037")

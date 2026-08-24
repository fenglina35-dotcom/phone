from pathlib import Path

from docx import Document


ROOT = Path(__file__).resolve().parents[1]
DOCS = ROOT / "docs/maintenance"
MARKER = "v1057／1.0.180 共同生活、宠物、电量与微信转账整合"


def append_section(filename, title, paragraphs):
    path = DOCS / filename
    document = Document(path)
    if any(MARKER in paragraph.text for paragraph in document.paragraphs):
        print(f"SKIP={path.name}")
        return
    document.add_heading(title, level=1)
    for paragraph in paragraphs:
        document.add_paragraph(paragraph)
    document.save(path)
    print(f"UPDATED={path.name}")


append_section(
    "AI开发项目_项目说明文档.docx",
    "v1057／1.0.180 共同生活、宠物、电量与微信转账整合（2026-08-24）",
    [
        "网页核心升级为 v1057，私人 iOS 升级为 1.0.180（180），原生桥继续为 25。共同生活入口和开关使用真实 button，并通过 cohabActionTap 消化同一次触摸产生的重复点击；cohabEnter 返回成功状态，仍进入原共同生活页面并保留原角色与生活记录。",
        "电子宠物在拖拽期间和结束后的短窗口内阻止浏览器合成点击，避免从食盆拖出食物投喂后误弹购买页；真正点击食盆购买和程序化投喂保持原行为。",
        "网页电量不再伪造 88%。支持 Battery API 的浏览器会绑定 BatteryManager，并在 pageshow、focus 和页面重新可见时刷新；iPhone Safari 不提供 Battery API，因此只能保留最近真实值或显示 --%，不能承诺读取系统真实电量。",
        "微信转账新增待收、已收、已退的独立明细页，记录转账、收款和退款时间，收款与退还幂等执行；旧 tcollect／treject 回执继续隐藏，黑白主题和返回聊天位置保持一致。",
        "根网页与私人 PhoneWeb.bundle 的 app.js、glass-theme.css、pet-game.js 和入口页面已按当前工作树同步。保留真实角色回复、朋友圈、后台通知、远控字幕、外置语音 Base URL／Key／模型／音色、私人云备份与真实外卖授权边界。Windows 完成语法与全量 1006/1006 自动测试；Mac 编译、签名和真实 iPhone 覆盖安装仍待执行。",
    ],
)

append_section(
    "AI开发项目_Bug记录模板.docx",
    "v1057／1.0.180 共同生活、宠物、电量与微信转账整合 Bug 记录（2026-08-24）",
    [
        "现象：共同生活入口偶发点击后不进入或重复触发；宠物从食盆拖拽后会误弹购买页；支持网页电量的浏览器恢复前台后不刷新且旧实现会伪造 88%；微信转账缺少完整明细、状态时间和一致的黑白主题。",
        "根因：共同生活用可点击容器并同时承接合成点击，缺少单次动作闸门；宠物 pointerup 后浏览器仍派发 scene click；BatteryManager 只在首次启动绑定，未处理页面恢复且显示层使用硬编码兜底；转账卡片仍沿用直接收款和历史回执结构。",
        "修复：增加 cohabActionTap 去重并让 cohabEnter 返回结果；增加宠物场景点击阻断窗口；重绑并刷新 BatteryManager，移除 88% 伪值；增加 transferDetail、幂等收款／退还、时间字段、存档指纹和主题样式，并隐藏旧重复回执。根网页改动同步到私人 Bundle。",
        "失败尝试与收口：版本升级后的首轮全量为 1004/1006，失败集中在 real-delivery-learning 与 real-delivery-role-handoff。排查确认不是外卖逻辑退化，而是并行微信转账收尾后根 app.js 有两行图标／反馈调整未同步到私人 Bundle；补齐后两项专项恢复通过。",
        "验证：根目录与私人 Bundle 的 app.js、pet-game.js 语法通过；共同生活、宠物、电量、转账及外卖交接专项通过；最终全量自动测试 1006/1006，git diff --check 与 ZIP 结构检查作为发布闸门。Windows 未执行 Xcode 编译、签名或真实 iPhone 验证。",
    ],
)

append_section(
    "AI开发项目_Bug修改规范.docx",
    "v1057／1.0.180 共同生活、宠物、电量与微信转账整合规范（2026-08-24）",
    [
        "触摸与点击规范：拖拽、长按或 pointer 捕获结束后，必须考虑浏览器额外派发的 click。对同一动作使用短时去重闸门；闸门只吞掉合成重复动作，真正点击和程序化调用必须保留。交互入口优先使用真实 button，并补齐 type、aria-label 和默认边框／内边距复位。",
        "网页电量规范：只有 navigator.getBattery 可用时才能声明网页读取到真实电量。BatteryManager 必须在 pageshow、focus 和 visibilitychange 恢复时刷新，并避免重复监听。不支持 Battery API 的 Safari 不得用固定百分比冒充真实值。私人原生桥电量与网页 Battery API 是两条独立能力，文案和测试不得混淆。",
        "转账状态规范：pending、received、refunded 必须由同一消息对象和幂等字段驱动；存档指纹需覆盖状态、时间与钱包结算标记。卡片点击只进入明细，不能直接改变余额；收款或退还只允许执行一次，历史辅助回执不得重复渲染。",
        "并行同步规范：当根网页仍可能被另一项任务收尾时，不能把旧 Bundle 反向覆盖根文件。每次发布前重新做根／私人逐行比较；任何晚到的根网页有效改动都要精确同步，然后重新运行受影响专项和全量测试。",
        "发布规范：版本号、Service Worker BUILD／HOTFIX、入口资源查询参数、PhoneWeb.bundle Info.plist、Xcode build／marketing version、原生 build marker、测试断言和安装说明必须同批更新。MacReady ZIP 只表示源码结构已检查；未在 Mac 编译签名和真机安装时必须明确写未验证。",
    ],
)

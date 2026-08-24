from pathlib import Path

from docx import Document


ROOT = Path(__file__).resolve().parents[1]
DOCS = ROOT / "docs/maintenance"
MARKER = "v1061／1.0.184"


def append_section(filename, title, paragraphs):
    path = DOCS / filename
    document = Document(path)
    matches = [index for index, paragraph in enumerate(document.paragraphs) if MARKER in paragraph.text]
    if len(matches) > 1:
        for paragraph in list(document.paragraphs[matches[1]:]):
            paragraph._element.getparent().remove(paragraph._element)
        document.save(path)
        print(f"DEDUPED={path.name}")
        return
    if matches:
        print(f"SKIP={path.name}")
        return
    document.add_heading(title, level=1)
    for paragraph in paragraphs:
        document.add_paragraph(paragraph)
    document.save(path)
    print(f"UPDATED={path.name}")


append_section(
    "AI开发项目_项目说明文档.docx",
    "v1061／1.0.184 外卖同意直传与角色自主发起修复（2026-08-25）",
    [
        "网页核心升级为 v1061，私人 iOS 升级为 1.0.184（184），原生桥继续为 25。用户明确提出点单后，角色当前真实模型回合若自然同意但漏掉结构化动作，系统只允许一次隐藏的真实模型动作补判；确认同意后立即把同一回合意图交给外卖自动化。",
        "已有起送价澄清任务优先沿用原 taskId、门店、主商品和进度。角色可以基于当前人设与上下文自主发起，但必须由当前真实模型明确输出结构化动作；普通聊天关键词、旧消息、页面重绘和后台恢复不能创建任务。",
        "普通角色回复、朋友圈回复、后台通知、远控字幕、共同生活、转账与外置语音线路保持不变。支付与平台安全验证仍由用户本人完成。",
        "根网页与私人 PhoneWeb.bundle 同步。Windows 完成语法、全量自动测试和 ZIP 结构检查；Mac 编译、签名和真实 iPhone 覆盖安装仍待执行。",
    ],
)

append_section(
    "AI开发项目_Bug记录模板.docx",
    "v1061／1.0.184 外卖同意直传与角色自主发起 Bug 记录（2026-08-25）",
    [
        "现象：用户首次明确说想吃什么，角色回复好、行、可以或等着，但真实模型没有同时输出外卖结构化动作，自动化没有启动；用户第二次重复要求后才可能启动。",
        "根因：可见自然回复与结构化动作是两条输出，安全闸门不能把固定同意词直接当成授权；但原流程在当前合法回合缺少一次受限的真实模型动作补判。",
        "修复：同一当前用户消息、账号、会话和角色回合内，只有明确点单语境加角色明确同意才允许一次隐藏补判。补判仍调用当前真实角色模型，只提取唯一结构化动作，不制造第二句固定角色回复。",
        "验证：覆盖首次同意直传、已有澄清优先续接、同消息幂等、普通聊天加好不触发、当前角色自主结构化动作可触发、旧消息与页面恢复不得启动。",
    ],
)

append_section(
    "AI开发项目_Bug修改规范.docx",
    "v1061／1.0.184 外卖同意直传与角色自主发起规范（2026-08-25）",
    [
        "自然回复规范：角色可见回复必须继续来自真实模型。结构化动作缺失时最多补判一次，补判结果只用于动作，不显示固定或伪造台词。",
        "授权规范：新任务必须绑定当前角色、账号、会话、消息与回合；角色自主点单必须是当前真实模型明确输出的结构化主动关心动作，不能由饿了、没吃饭等关键词直接启动。",
        "续接规范：已有澄清先于新任务补判，回答必须沿用原 taskId 和原进度。同一 messageId 重复调用不得二次启动、二次搜索或重复加购。",
        "发布规范：网页、私人 Bundle、测试、版本号、维护文档和安装包同批更新。未做 Mac 编译、签名和真机验收时必须如实标注。",
    ],
)

"""Append the v1081 natural-delivery and KFC homepage release to maintenance records."""

from pathlib import Path

from docx import Document
from docx.enum.text import WD_BREAK


ROOT = Path(__file__).resolve().parents[1]
DOCS = ROOT / "docs" / "maintenance"
MARKER = "v1081｜"


def append_docx(filename: str, title: str, paragraphs: list[str]) -> None:
    path = DOCS / filename
    doc = Document(path)
    existing = "\n".join(p.text for p in doc.paragraphs)
    if MARKER in existing:
        raise RuntimeError(f"{filename}: v1081 section already exists")
    page = doc.add_paragraph()
    page.add_run().add_break(WD_BREAK.PAGE)
    heading = doc.add_paragraph()
    run = heading.add_run(title)
    run.bold = True
    run.font.size = doc.styles["Normal"].font.size
    for paragraph in paragraphs:
        doc.add_paragraph(paragraph)
    doc.save(path)


def append_txt(filename: str, title: str, paragraphs: list[str]) -> None:
    path = DOCS / filename
    existing = path.read_text(encoding="utf-8")
    if MARKER in existing:
        raise RuntimeError(f"{filename}: v1081 section already exists")
    block = "\n\n" + title + "\n" + "\n".join(paragraphs) + "\n"
    path.write_text(existing.rstrip() + block, encoding="utf-8", newline="")


records = [
    (
        "AI开发项目_项目说明文档",
        "v1081｜自然点单解析与 KFC 首页套餐（2026-08-27）",
        [
            "版本：共享网页 v1081；私人 iOS 1.0.206（206）；原生桥 25。v1080 的真实外卖角色第一人称回执、付款事实边界以及此前修复继续保留。",
            "当前角色点单回合中，通用“品牌的商品，规格”自然句可以直接形成门店、商品和规格，不再强制用户说“点”、加冒号或使用机器式字段。例如“百分茶的暴打土芭乐柠檬茶，不另加糖”会规范为门店百分茶、商品暴打土芭乐柠檬茶、规格不另外加糖；模型动作残缺时也从同一句补全。该能力不是百分茶白名单。",
            "解析器在启动前统一拦截历史陈述、明确拒绝和评价询问。类似“我以前喝过…”“不要点…”“你觉得…怎么样”不会创建真实外卖任务，也不会进入隐藏补判。",
            "KFC 当前首页四件套按真实名称“【夜宵专享】吃堡堡4件套”识别；首页已经存在套餐时不再店内搜索套餐名，泛化任务说明也不会被误当成独立加购。九珍果汁饮料、爆汁三柠茶的真实“选规格”步骤仍待门店营业后真人环境验证，不能记录为已完成。",
        ],
    ),
    (
        "AI开发项目_Bug修改规范",
        "v1081｜自然外卖意图与首页套餐边界规范（2026-08-27）",
        [
            "自然外卖解析应支持通用“品牌的商品，规格”结构，不得靠单一品牌白名单；只有当前角色点单回合、商品具有食物语义且尾部规格属于受限白名单时才能直接启动或补全残缺动作。",
            "历史陈述、否定/取消和评价询问必须在显式解析与隐藏补判之前共同拦截。不能因为句子含品牌、饮品或食物名就创建真实任务。",
            "首页限定套餐必须优先使用首页真实卡片，不得在店内搜索套餐名。工作流说明、禁止重复单点等文字不能被解析成独立加购；只有套餐未覆盖的明确食品项才允许继续搜索。",
            "自动测试通过只能证明解析和页面状态机约束。门店打烊时未执行的真实规格选择必须写成待验收，禁止把九珍果汁饮料或爆汁三柠茶的“选规格”写成已完成。",
        ],
    ),
    (
        "AI开发项目_Bug记录模板",
        "v1081｜自然品牌商品句未启动与 KFC 套餐误搜索（2026-08-27）",
        [
            "现象：用户自然说“品牌的商品，规格”时，因没有“点”或冒号可能无法在当前回合启动；模型结构化动作残缺后也未能从原句补全。KFC 首页套餐名称变化后，旧逻辑可能搜索旧套餐名，或把任务说明误当独立商品。",
            "根因：显式意图过度依赖动作词和标签格式，缺少受控的通用品牌-商品-规格解析；隐藏补判没有共享同一误触拦截。KFC 默认套餐写死旧称，首页限定条件与店内搜索/独立加购过滤不完整。",
            "修复：新增上下文自然点单解析与统一拒绝门禁；尾部规格仅接受受限词集并标准化“不另加糖”。KFC 改用当前首页四件套签名识别，首页限定任务禁止店内搜套餐，并过滤泛化说明。",
            "验证：根目录全量、自然澄清专项和外卖服务全量测试在发布前重新执行；根/私人 delivery.js 必须逐字节一致。没有发起新的真实订单。KFC 九珍/三柠真实规格步骤仍待营业后验证。",
        ],
    ),
    (
        "AI开发项目_新聊天启动说明",
        "v1081｜新聊天接手说明（2026-08-27）",
        [
            "当前候选：网页 v1081；私人 iOS 1.0.206（206）；原生桥 25。先核对远端提交、ZIP SHA-256、根/私人 delivery.js 哈希和大量未跟踪资料保护状态。",
            "自然句“品牌的商品，规格”在当前点单回合可以直接启动并补全模型残缺动作；这是一条通用解析规则，不是百分茶特例。历史陈述、明确拒绝和评价询问不能触发。",
            "KFC 首页默认套餐当前识别“【夜宵专享】吃堡堡4件套”，首页已有套餐不得店内搜索。九珍果汁饮料/爆汁三柠茶真实规格选择尚未在营业门店验证，必须继续标为待验收。",
            "Windows 自动测试和 ZIP 结构校验不能写成 Mac 编译、签名、真实 iPhone 或营业中 KFC 已通过；未跟踪诊断图和用户资料不得清理。",
        ],
    ),
]

for stem, title, paragraphs in records:
    append_docx(stem + ".docx", title, paragraphs)
    append_txt(stem + ".txt", title, paragraphs)

print("Updated four maintenance DOCX/TXT pairs for v1081")

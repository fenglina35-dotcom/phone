from pathlib import Path

from docx import Document


ROOT = Path(__file__).resolve().parents[1]
DOCS = ROOT / "docs" / "maintenance"


def append_section(filename, title, paragraphs):
    path = DOCS / filename
    doc = Document(path)
    doc.add_heading(title, level=1)
    for text in paragraphs:
        doc.add_paragraph(text)
    doc.save(path)


append_section(
    "AI开发项目_项目说明文档.docx",
    "当前发布基线｜网页 v1011／私人 iOS 1.0.132 (132)／原生桥 25（2026-08-21）",
    [
        "本版对三条互相独立的故障链做最小修复。AI 账户迁往邀请码项目后，phone-ai 地址与公开项目 Key 必须成对选择；外置 Mossland／Fish 仍保留用户自己的 base、Key 和音色，只修正它们经过 phone-ai 中转时的 Supabase 网关鉴权。",
        "朋友圈评论继续只接受真实模型输出，不生成固定或伪造回复。它现在与普通微信一样读取角色主／副模型选择，并把最近六轮已标明说话人的微信上下文、当前朋友圈与评论线程交给同一条真实线路；提示词规模恢复为有界范围。",
        "后台立即测试与正式主动联系优先使用角色当前外置模型线路，并同步另一条已配置外置线路作为备用。显式测试每条线路只尝试一次，失败后结束并报告真实 provider／APNs 结果，不再隔一分钟重复占用普通微信。phone-role-push 已部署为独立云 ACTIVE v9。",
        "Windows 专项和完整回归、PhoneWeb.bundle 重建、ZIP 结构与校验完成后才可交付。Mac 编译、Apple 签名以及真实 iPhone 上的朋友圈回复、后台通知到达与入聊、外置语音播放仍须真机验收，Windows 结果不能代替这些结论。",
    ],
)

append_section(
    "AI开发项目_Bug记录模板.docx",
    "2026-08-21｜v1011｜AI 迁移 401、朋友圈固定主模型、后台单线路超时",
    [
        "现象：网页与私人 App 的角色语音同时突然报 HTTP 401；普通微信能回复而朋友圈评论持续失败；立即模拟后台主动消息长期显示调用当前模型，最终没有 APNs 通知，并会拖慢普通微信。",
        "语音根因：AI_BACKEND_URL 已迁到 lovbzibismsjqvjujilz，但 aiRelay 请求头仍固定使用旧 lkhlyfpssmrjkkzhuzag 的 GATE_KEY。真实网关对照为新地址＋新 Key 返回 200，新地址＋旧 Key 返回 401 Invalid API key。外置 Mossland／Fish 的 external_tts 也经过 aiRelay，所以被同一网关错误连带影响；不是用户外置平台 Key 或余额失效。",
        "朋友圈根因：普通微信请求明确传 aux:c.model==='aux'，朋友圈自 7 月 30 日以来一直漏传角色模型选择，固定走主模型；当角色实际选择副模型或主线路后来不可用时，表现为只有朋友圈失败。v1003 加入真实失败状态后问题变得可见，v1008 扩大的上下文又增加了超时概率。后台测试早于或晚于朋友圈操作都不改变此独立根因。",
        "后台根因：独立云任务日志明确记录 one_minute_test 失败在 profile-current:timeout，未生成 outbox，因而根本未到 APNs。APNs 令牌与其他历史 outbox 已有 Apple 接受记录；不能把此故障误判为通知权限。独立云没有内置模型 Secrets，服务器只同步一条当前外置线路，18 秒超时后又隔 60 秒重跑，形成长时间占用。",
        "修复：aiCoreKey(url) 按目标项目 origin 选择对应公开 Key；朋友圈传入角色实际主／副模型选择并缩减上下文；后台同步当前与备用两条外置线路，单线路 27 秒、有界总时限，显式测试不再跨分钟重试。全程不修改用户外置平台配置，不生成假角色回复。",
        "验证：新 phone-ai 地址＋新 Key 在线返回 200，新地址＋旧 Key 返回 401；专项测试覆盖地址／Key 同项目、朋友圈副模型、真实失败不伪造、后台双线路与单轮释放。phone-role-push 已部署 ACTIVE v9。仍需新安装包在真实 iPhone 验证三项最终用户结果。",
    ],
)

append_section(
    "AI开发项目_Bug修改规范.docx",
    "v1011 新增规范：项目鉴权、场景模型路线与后台生成释放（2026-08-21）",
    [
        "Supabase Edge Function 地址和 apikey／Authorization 必须属于同一项目。迁移 endpoint 时禁止只替换 URL 或把所有兼容地址一律硬换成新 Key；应通过唯一的 endpoint→Key 选择器处理，并以新地址＋新 Key、新地址＋旧 Key、旧兼容映射三组自动测试锁定。",
        "所有角色可见生成场景必须继承该角色实际主／副模型选择。普通微信、朋友圈、电话、共同生活或后台任务若存在独立调用点，必须逐一核对 aux／routeIndex，而不能因共用 chatAPI 就假设线路相同。",
        "朋友圈失败不得通过固定句、旧回复复制或本地伪造掩盖。应保留真实错误，模型已有真实可见输出仅因兼容解析误清空时可以回收原始可见正文，但 JSON／控制外壳不得直接展示。",
        "显式后台测试失败应在一次有界任务内尝试已同步的外置线路并立即释放；不得用分钟级重复任务长期占用前台同一模型。入队、模型生成、outbox 持久化、APNs 接受和客户端入聊必须分别记录，任何前一层成功都不能冒充后一层送达。",
    ],
)

append_section(
    "AI开发项目_新聊天启动说明.docx",
    "2026-08-21 v1011 接手补充",
    [
        "当前候选基线为网页 v1011、私人 iOS 1.0.132 (132)、原生桥 25。AI 账户 phone-ai 位于邀请码项目 lovbzibismsjqvjujilz；真实 iPhone 伴生、后台任务和 APNs 仍只位于独立云 qvuahlqimcfgeoetosnl，禁止混用或覆盖。",
        "本版三条根因已分离：语音 401 是新 AI 地址携带旧项目 Key；朋友圈失败是漏传角色主／副模型选择；立即后台测试失败发生在 profile-current 模型超时、尚未到 APNs。后续真机验收必须分别检查，不得再用一个现象解释另外两个。",
        "接手后先只读核对 Git、独立云 phone-role-push ACTIVE v9、唯一 PhoneWeb.bundle 和 v1011／1.0.132 版本一致性。Mac 编译及真实 iPhone 验收未完成前，不得写成已通过。",
    ],
)

print("Updated four maintenance documents for v1011")

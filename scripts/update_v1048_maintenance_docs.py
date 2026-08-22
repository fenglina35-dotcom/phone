from pathlib import Path

from docx import Document
from docx.enum.text import WD_BREAK


ROOT = Path(__file__).resolve().parents[1]
DOCS = ROOT / "docs" / "maintenance"
MARKER = "v1048｜"


def append_section(filename, title, paragraphs):
    path = DOCS / filename
    doc = Document(path)
    existing = "\n".join(p.text for p in doc.paragraphs)
    if MARKER in existing:
        raise RuntimeError(f"{filename}: v1048 section already exists")
    page = doc.add_paragraph()
    page.add_run().add_break(WD_BREAK.PAGE)
    heading = doc.add_paragraph()
    run = heading.add_run(title)
    run.bold = True
    run.font.size = doc.styles["Normal"].font.size
    for paragraph in paragraphs:
        doc.add_paragraph(paragraph)
    doc.save(path)


append_section(
    "AI开发项目_项目说明文档.docx",
    "v1048｜私人主设备同步、真实外卖换店与微信头像整合（2026-08-23）",
    [
        "版本：共享网页 v1048；私人 iOS 1.0.166（166）；原生桥 25。以 origin/main 的 v1047 为基线，网页核心重新同步进唯一的私人 PhoneWeb.bundle。",
        "云同步改为私人 App 主设备、网页版镜像：私人 App 保存并完成防抖备份后，以 source=private-ios、单调递增 revision、capturedAt、updatedAt 和 sourceBuild 写入同一个云ID记录；网页版启动、回前台或重新聚焦时只读取更高的私人修订。",
        "网页不再把旧快照自动写回私人记录。首次绑定或网页存在本地改动时不静默覆盖：先提示冲突，默认按钮保留私人版本；应用私人镜像前把网页完整状态保存为可撤回快照。并发写采用 updated_at 条件更新，远端已变化时拒绝本次覆盖。旧云ID的手动备份与恢复入口仍兼容。",
        "常规业务状态包括角色、聊天、记忆、朋友圈、钱包与设置。放映室原生大视频、系统相册原文件等大型媒体不属于网页镜像范围，界面已明确提示，不再宣称全部媒体都已同步。",
        "真实外卖保持可见浏览器和人工支付边界：优先复用已验证路线，在店内精确搜索目标商品；第一家打烊时最多继续检查三家，并把打烊/未完成结果作为系统提示，不伪装成角色台词。平台人机验证不能绕过，出现时暂停同一任务等待本人完成后继续，超时进入冷却，避免无限循环。",
        "微信聊天双方头像统一为 42px 正方形圆角并与气泡垂直居中。只调整头像尺寸和对齐，不改变普通角色回复、朋友圈回复、后台通知、远程控制或外置语音 base/key/voice/model 配置。",
    ],
)

append_section(
    "AI开发项目_Bug记录模板.docx",
    "v1048｜网页云恢复反向覆盖私人版本与外卖打烊循环（2026-08-23）",
    [
        "现象一：私人手机号备份和网页 Supabase backups 是两套独立快照；网页 cloudRestore 会整份替换本机状态，网页自动备份还可能把旧网页数据写回云端，用户无法保证私人版本是主数据源。",
        "根因一：旧记录没有来源、修订号和并发条件，只有 updated_at；自动备份默认双向覆盖，首次绑定与并发修改没有冲突门槛，也没有恢复前回滚快照。",
        "处理一：新增 private-primary 镜像元数据、单调 revision、来源与采集时间；私人端条件写入，网页端只拉取更高修订。首次绑定或网页本地变更弹出冲突确认，应用前保存 IndexedDB 回滚快照；网页自动备份发现私人主记录时立即停止。",
        "现象二：真实外卖第一家店打烊或平台出现图片验证时，角色可能反复说等待、任务无限搜索，用户看不到明确结束原因。",
        "处理二：已验证路线优先、最多三家、店内精确商品搜索；打烊自动换下一家，全部打烊立即结束并通知；验证码只暂停当前任务并等待人工完成，超时后设置冷却，不用规避检测的手段绕过平台验证。",
        "验证：两份 app.js 语法通过；云同步/外卖/头像定向 38/38 通过；网页与私人 Bundle 的云同步实现逐字一致。完整 Node 测试、Mac 编译签名和真机云端/外卖联调在发布前继续执行并记录。",
    ],
)

append_section(
    "AI开发项目_Bug修改规范.docx",
    "v1048｜主设备镜像与平台验证处理规范（2026-08-23）",
    [
        "多端整份状态同步必须明确主设备。快照至少携带 source、revision、capturedAt/updatedAt 和版本；revision 只增不减。旧端不得自动覆盖主设备，新端首次绑定不得静默选择任一端。",
        "检测到并发改动时，先保留本机状态并明确提示；应用远端快照前必须保存可撤回的完整快照。写云端时使用服务端条件更新或等价并发保护，条件不匹配就停止，禁止最后写入者无条件获胜。",
        "状态同步文案必须如实区分业务状态和大型媒体。没有进入 fullBackupState 或超过云端容量的原生视频、系统媒体原文件，必须明确写为不在镜像范围内。",
        "浏览器自动化不得隐藏、绕过或规避平台验证码和风控。允许的处理是使用真实可见浏览器、复用已登录会话、降低无意义重复搜索、暂停当前操作等待本人验证、验证完成后继续同一任务，并在超时后有限冷却。",
        "真实外卖搜索必须有有界停止条件。已验证路线优先；门店打烊应尝试下一条匹配路线并通知最终结果；不得把第一家、第一项或模糊命中的商品当作用户指定商品，不得用角色台词掩盖系统失败。",
    ],
)

append_section(
    "AI开发项目_新聊天启动说明.docx",
    "v1048｜发布与私人主同步交接（2026-08-23）",
    [
        "当前发布目标：网页 v1048；私人 iOS 1.0.166（166）；原生桥 25；分支 main。v1048 必须包含 v1047 已上线内容、真实外卖换店/偏好、42px 聊天头像和私人主设备云镜像。",
        "私人 App 是云同步主设备。私人端与网页端填同一云地址、Key 和云ID；私人端保存后上传更高 revision，网页启动或回前台只读取更高的私人 revision。网页发现本地改动时先保留并提示，不会自动上传覆盖私人版本。",
        "旧云ID恢复仍可手动使用。首次绑定不自动覆盖；应用私人镜像前保存网页撤回快照。聊天、角色、记忆、朋友圈、钱包和设置会同步，原生大视频与系统媒体原文件不在网页镜像范围内。",
        "真实外卖仍只自动执行到支付前：平台真实验证码必须由本人完成，支付也由本人确认。第一家打烊会检查下一家，最多三家；全部打烊、金额不足或验证超时要明确结束，不能无限循环。",
        "发布前要求：全量 Node 测试通过；只有一个最新 MacReady ZIP；包内一个 PhoneWeb.bundle/index.html，无嵌套 ZIP、缓存、预览图和旧安装说明。Windows 复核不能替代 Mac 编译、签名与真机验证，未完成项必须如实报告。",
    ],
)

print("Updated four maintenance DOCX files for v1048")

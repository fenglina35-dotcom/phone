from pathlib import Path

from docx import Document


ROOT = Path(__file__).resolve().parents[1]
DOCS = ROOT / "docs" / "maintenance"


def append_section(path: Path, heading: str, paragraphs: list[str]) -> None:
    doc = Document(path)
    if any(p.text.strip() == heading for p in doc.paragraphs):
        raise RuntimeError(f"section already exists: {heading}")
    doc.add_heading(heading, level=1)
    for paragraph in paragraphs:
        doc.add_paragraph(paragraph)
    doc.save(path)


append_section(
    DOCS / "AI开发项目_项目说明文档.docx",
    "2026-09-06 私人共同生活键盘作用域修复候选",
    [
        "发布范围：仅私人 iOS 壳层，候选版本 v1186／1.0.312 (312)。公开网页源码、网页发布和私人 PhoneWeb.bundle 的 v1184 业务功能不变；只更新私人运行时诊断版本标识。",
        "问题复核：v1185 把所有页面共用的 WKWebView 放进 UIKeyboardLayoutGuide 容器，作用域从共同生活扩大到微信，真机出现输入栏滞留、底部黑区、页面弹跳和微信回归。因此 v1185 方案撤回，不再继续调节该全局约束。",
        "候选修复：恢复直接 WKWebView 和 SwiftUI 系统键盘安全区路径。只在共同生活输入框 #off_in 聚焦期间暂停 WKWebView 外层滚动，避免 WebKit 自动聚焦滚动与系统整页上推叠加；键盘隐藏、切换其他编辑框或页面销毁时恢复。禁止 KVO 监听 contentOffset、禁止 setContentOffset、禁止全局 viewport 或键盘容器改写。",
        "验证边界：Windows 运行作用域与全量回归测试并核对公开网页零改动；真实动画仍必须在 Mac/Xcode 编译签名后，于同一台 iPhone 分别重复验证共同生活首次打开、收起、旁白切换和微信打开／输入／收起。通过真机前只能标记为候选。",
    ],
)

append_section(
    DOCS / "AI开发项目_Bug记录模板.docx",
    "v1186 私人 1.0.312 Bug 记录：共同生活键盘作用域回退与隔离",
    [
        "用户可见现象：v1185 安装后共同生活输入栏在键盘收起后卡在上方并留下黑色空区，首次打开仍弹跳；原本正常的微信输入框也被同一改动带坏。该证据证明故障不是单一动画时长问题，而是共享原生容器受到了过宽改动。",
        "失败方案记录：v1183 冻结键盘安全区并操作滚动偏移，真机会卡顿和错位；v1184 恢复系统路径后明显改善但仍有局部双滚动；v1185 使用全局 UIKeyboardLayoutGuide，造成共同生活恶化并回归微信。三条路线不得继续叠加或微调后重试。",
        "根因判断：共同生活页面内已有自己的消息滚动区，WebKit 聚焦输入框时还会尝试移动外层 WKScrollView；如果原生壳同时改变全部页面的底部约束，就形成两个布局权威。v1185 将局部问题错误扩大到整个 WebView 生命周期。",
        "v1186 候选处理：撤回全局键盘导向容器；所有页面恢复直接 WKWebView。原生注入只识别 #off_in 的 focusin／focusout，在该焦点生命周期临时关闭外层滚动，并在 keyboardDidHide、切换其他输入框及 dismantle 时恢复。微信不进入此作用域。",
        "当前状态：源码静态检查、针对性测试和全量测试通过后可打包；Windows 无法编译 iOS 或观察真实键盘动画，必须由 Mac/Xcode 和真实 iPhone 覆盖安装后确认，未验收前不得写成已修复。",
    ],
)

append_section(
    DOCS / "AI开发项目_Bug修改规范.docx",
    "新增规则：私人键盘修复必须限定目标输入框并反向验证共享页面",
    [
        "任何键盘 Bug 修改前，先列出共享 WKWebView、SwiftUI 根视图、viewport、滚动容器和焦点状态的全部消费者；明确目标页、非目标页和恢复路径。不能用局部页面现象推导全局壳层都需要改动。",
        "只有共同生活输入框异常时，不得全局替换 WKWebView 容器约束，不得给根视图新增全局键盘忽略规则，不得监听并强制重写 contentOffset，也不得修改所有页面共用的 visualViewport 行为。优先以稳定基线恢复非目标页面，再对目标 DOM 输入框做可撤销、可验证的最小作用域处理。",
        "每个临时状态必须覆盖正常结束和异常结束：blur、切换到另一输入框、键盘交互式隐藏、页面退出和 WebView 销毁都要恢复。测试必须证明状态只由目标输入框开启，且不会泄漏到微信、通话或其他编辑框。",
        "私人包发布门槛新增双向回归：共同生活至少重复三轮打开／收起／旁白切换；微信至少重复一轮进入聊天、打开键盘、输入、收起和再次打开。Windows 代码测试只能证明约束和作用域，真实 iPhone 动画没有验收时必须明确标为 Mac 待编译、真机待确认。",
    ],
)

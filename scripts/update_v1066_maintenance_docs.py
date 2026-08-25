from pathlib import Path

from docx import Document


ROOT = Path(__file__).resolve().parents[1]
DOCS = ROOT / "docs/maintenance"
MARKER = "v1066／1.0.189"


def append_section(filename, title, paragraphs):
    path = DOCS / filename
    document = Document(path)
    if any(MARKER in paragraph.text for paragraph in document.paragraphs):
        print(f"SKIP={path.name}")
        return
    document.add_heading(title, level=1)
    for paragraph in paragraphs:
        document.add_paragraph(paragraph)
    document.save(path)
    print(f"UPDATED={path.name}")


append_section(
    "AI开发项目_项目说明文档.docx",
    "v1066／1.0.189 伴生轮询日期格式化卡顿发热修复（2026-08-25）",
    [
        "网页核心升级为 v1066，私人 iOS 升级为 1.0.189（189），原生桥继续为 25。本次只修复已有 Instruments 证据指向的私人 App WebContent 热点，不再通过猜测性放慢或关闭其他定时器。",
        "用户提供的官方 Instruments XML 显示，卡顿期间 WebContent 约 96% 的 CPU 样本位于 DOM 定时器路径，约 91% 位于 JSC::constructIntlDateTimeFormat；首次异常样本约在启动后 3.184 秒。代码中的伴生快照轮询在启动后 3.2 秒运行，原日期核对会每次新建 Intl.DateTimeFormat，时间和调用路径形成直接对应。",
        "私人 App 的屏幕时间报告与 WKWebView 处于同一 iPhone 系统时区，因此原生环境直接使用本机年月日，不再从伴生轮询定时器构造 Intl formatter；网页环境仍支持服务器报告的指定时区，但每个时区只构造一次并缓存复用。",
        "伴生轮询频率、数据语义、普通角色回复、朋友圈回复、后台通知、远控字幕、通话、内外置语音、共同相册与外卖链路不改。Windows 自动测试通过；Mac 编译、签名和本版本真实 iPhone 长时间验证仍待完成。",
    ],
)

append_section(
    "AI开发项目_Bug记录模板.docx",
    "v1066／1.0.189 私人 App 偶发卡顿发热 Bug 记录（2026-08-25）",
    [
        "现象：私人 App 可连续顺畅约数日后再次出现偶发严重卡顿；滑动可能仍可进行，但点击响应很慢，打开小应用偶发白屏并回到主屏，手机温度快速上升，后台消息也可能同期停止显示。顺畅恢复后温度随之下降。",
        "已排除：这不是单凭截图可归因的手机硬件故障；多轮广泛调整动画、图片恢复、定位或普通定时器没有产生稳定改善。原始 trace 在 Windows 直接解析地址会受 Instruments 压缩和重定位影响，不能据此猜函数；本次只采用 Mac 官方导出的 XML 调用栈。",
        "证据：官方导出中 WebContent 的绝大多数 CPU 样本落在 DOM 定时器和 Intl.DateTimeFormat 构造路径，第一批样本与启动后 3.2 秒的伴生轮询对齐。",
        "修复：私人 App 的 companionUsageDayAt 使用本机日期快速路径；网页指定时区按时区缓存 formatter。新增回归测试确保私人轮询即使 Intl 构造器抛错也可完成，并确保网页同一时区只构造一次。",
        "未完成验证：Windows 自动化不能替代 Xcode 编译和真机热状态观察。安装 v1066 后需从冷启动观察至少 2 分钟，并继续正常使用数日；若复现，应以新 trace 和准确时间继续定位，而不是重复改同一处。",
    ],
)

append_section(
    "AI开发项目_Bug修改规范.docx",
    "v1066／1.0.189 WebContent 性能取证与修复规范（2026-08-25）",
    [
        "取证规范：偶发卡顿和发热必须优先取得 Instruments Time Profiler、Hangs 与 Thermal State；使用官方导出的表格或已符号化调用栈。未经符号化的原始地址不得直接映射为项目函数。",
        "因果规范：修改前必须同时给出热点调用链、首次出现时间与代码调度时间。只有三者吻合才进入修复；没有新变量时禁止重复放慢全部定时器、关闭后台能力或重写稳定链路。",
        "实现规范：高频或延迟定时器内不得重复构造 Intl formatter、正则、解析器或大对象；可由同设备本地日期完成时走无 Intl 快速路径，确需时区格式化时按稳定键缓存。不得用旧快照或伪数据冒充实时结果。",
        "保护规范：性能修复不得改变真实模型回复、朋友圈、APNs 与消息落库、远控字幕、伴生真实数据、通话、外置语音配置和共同相册。系统进度或固定文本不得冒充角色回复。",
        "发布规范：自动测试只能证明源码回归边界；MacReady 不等于已编译或真机已修好。版本、Bundle、Xcode build、安装说明、维护记录和 ZIP 必须一致，真机结果需单独记录。",
    ],
)
